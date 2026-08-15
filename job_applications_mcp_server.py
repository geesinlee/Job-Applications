#!/usr/bin/env python3
"""Job Applications MCP Server — Orchestrates job application workflow.

Manages company folders, job description parsing, research templates,
territory mapping, profile management, match scoring, gap analysis,
application tracking, and document generation for job applications.

Deployment:
  - stdio  : Mac (transient, during active Claude session)
  - http   : pi-4 gs-pi-4 :8086 (always-on systemd service)

Artefact files (JD.md, CVs, research) live on the NAS share.
tracker.json and profile.json live on pi-4 local disk, rsynced to NAS
after every write.
"""

import difflib
import json
import os
import re
import shutil
import subprocess
import sys
import uuid
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # python-dotenv optional; env vars may be set by systemd EnvironmentFile

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

import requests
from bs4 import BeautifulSoup

from mcp.server.fastmcp import FastMCP
from mcp.server.auth.provider import AccessToken, TokenVerifier
from mcp.server.auth.settings import AuthSettings

from requirement_service import RequirementService
from cv_versioning_service import CVVersioningService

__version__ = "0.3.0"

# ---------------------------------------------------------------------------
# Configuration — env vars with __file__-relative fallbacks
# ---------------------------------------------------------------------------

_SRC_DIR = Path(__file__).resolve().parent

BASE_DIR = Path(os.environ.get("JOB_APP_BASE_DIR", str(_SRC_DIR)))
ARTEFACTS_DIR = Path(os.environ.get("JOB_APP_ARTEFACTS_DIR", str(BASE_DIR)))
TRACKER_PATH = Path(os.environ.get("JOB_APP_TRACKER_PATH", str(BASE_DIR / "tracker.json")))
PROFILE_PATH = Path(os.environ.get("JOB_APP_PROFILE_PATH", str(BASE_DIR / "profile.json")))
BASE_CV_PATH = Path(os.environ.get(
    "JOB_APP_BASE_CV_PATH",
    str(ARTEFACTS_DIR / "DXC" / "CV LEE Gee Sin 2026 - DXC Client Partner Public Sector.md"),
))
NAS_SYNC_PATH = os.environ.get("NAS_SYNC_PATH", "")   # e.g. gs@rv-cloud.local:/share/job-app-data/
MCP_MODE = os.environ.get("MCP_MODE", "stdio")         # "stdio" | "http"
MCP_AUTH_TOKEN = os.environ.get("MCP_AUTH_TOKEN", "")

# ---------------------------------------------------------------------------
# Startup validation
# ---------------------------------------------------------------------------

def _startup_validate() -> None:
    """Verify BASE_DIR exists; initialise tracker.json and profile.json if absent.
    Exits with code 1 (writing to stderr) if BASE_DIR is missing.
    """
    if not BASE_DIR.is_dir():
        sys.stderr.write(
            f"[job-applications] ERROR: BASE_DIR not found or not a directory: {BASE_DIR}\n"
            "Set JOB_APP_BASE_DIR to a valid path.\n"
        )
        sys.exit(1)

    if not TRACKER_PATH.exists():
        TRACKER_PATH.parent.mkdir(parents=True, exist_ok=True)
        TRACKER_PATH.write_text(
            json.dumps({"schema_version": "1.0", "applications": []}, indent=2),
            encoding="utf-8",
        )

    if not PROFILE_PATH.exists():
        PROFILE_PATH.parent.mkdir(parents=True, exist_ok=True)
        PROFILE_PATH.write_text(
            json.dumps({"schema_version": "1.0"}, indent=2),
            encoding="utf-8",
        )

_startup_validate()

# ---------------------------------------------------------------------------
# Persistence helpers
# ---------------------------------------------------------------------------

def _utc_now() -> str:
    """Return current UTC time as ISO-8601 string, e.g. '2026-08-02T14:30:00Z'."""
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _load_tracker() -> dict:
    """Load tracker.json, returning empty schema on any read error."""
    try:
        return json.loads(TRACKER_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {"schema_version": "1.0", "applications": []}


def _save_tracker(data: dict) -> None:
    """Save tracker.json and rsync to NAS backup (fire-and-forget)."""
    TRACKER_PATH.write_text(
        json.dumps(data, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    _nas_sync()


def _load_profile() -> dict:
    """Load profile.json, returning empty schema on any read error."""
    try:
        return json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {"schema_version": "1.0"}


def _save_profile(data: dict) -> None:
    """Save profile.json and rsync to NAS backup (fire-and-forget)."""
    PROFILE_PATH.write_text(
        json.dumps(data, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    _nas_sync()


def _nas_sync() -> None:
    """Rsync tracker.json and profile.json to NAS backup path (non-blocking).
    Silently skips if NAS_SYNC_PATH is not configured.
    """
    if not NAS_SYNC_PATH:
        return
    try:
        subprocess.Popen(
            ["rsync", "-a", str(TRACKER_PATH), str(PROFILE_PATH), NAS_SYNC_PATH],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
    except FileNotFoundError:
        pass  # rsync not installed; skip silently


def _get_or_create_evidence_service():
    """Lazy-load EvidenceService instance (temporary integration point).

    In Gate 6+, this will be replaced with proper dependency injection.
    For now, we create a minimal instance pointing to the tracker/profile store.
    """
    # Placeholder: returns a stub that queries tracker evidence
    # This will be replaced when Gate 4 (EvidenceService) is integrated
    # For now, score_match continues to work via Claude's LLM-based scoring
    # and extracted_requirements are provided for Claude's reference

    class TemporaryEvidenceStub:
        def query_evidence(self, **filters):
            # Stub: will be replaced when Gate 4 is integrated
            return []

    return TemporaryEvidenceStub()


_cv_service_instance = None


def _get_or_create_cv_service():
    """Get or create CVVersioningService singleton.

    Initializes with RequirementService and temporary evidence stub.
    """
    global _cv_service_instance
    if _cv_service_instance is None:
        evidence_service = _get_or_create_evidence_service()
        requirement_service = RequirementService(evidence_service)
        _cv_service_instance = CVVersioningService(requirement_service, evidence_service)
    return _cv_service_instance


# ---------------------------------------------------------------------------
# Stage machine constants
# ---------------------------------------------------------------------------

VALID_STAGES = {
    "new", "applied", "screening",
    "interview_r1", "interview_r2", "interview_r3",
    "offer", "accepted", "rejected", "withdrawn",
}
TERMINAL_STAGES = {"accepted", "rejected", "withdrawn"}
INTERVIEW_STAGES = {"interview_r1", "interview_r2", "interview_r3"}

VALID_TRANSITIONS: dict[str, set[str]] = {
    "new":          {"applied", "rejected", "withdrawn"},
    "applied":      {"screening", "rejected", "withdrawn"},
    "screening":    {"interview_r1", "rejected", "withdrawn"},
    "interview_r1": {"interview_r2", "offer", "rejected", "withdrawn"},
    "interview_r2": {"interview_r3", "offer", "rejected", "withdrawn"},
    "interview_r3": {"offer", "rejected", "withdrawn"},
    "offer":        {"accepted", "rejected", "withdrawn"},
    "accepted":     set(),
    "rejected":     set(),
    "withdrawn":    set(),
}

# ---------------------------------------------------------------------------
# Path resolver helpers
# ---------------------------------------------------------------------------

def _make_role_slug(role_title: str) -> str:
    """Convert role title to URL-safe lowercase slug.
    e.g. 'Enterprise AE – Strategic Accounts' -> 'enterprise-ae-strategic-accounts'
    """
    slug = role_title.lower()
    slug = re.sub(r"[^\w\s-]", "", slug)   # remove special chars except hyphen
    slug = re.sub(r"[\s_]+", "-", slug)    # spaces/underscores to hyphens
    slug = re.sub(r"-+", "-", slug)        # collapse multiple hyphens
    return slug.strip("-")


class AmbiguousRoleError(ValueError):
    """Raised when a company has multiple roles and none is specified."""
    def __init__(self, company: str, roles: list[str]):
        self.company = company
        self.roles = roles
        super().__init__(f"Multiple roles at {company}: {roles}")


def _resolve_company_folder(
    company: str,
    role_title: str | None = None,
    tracker: dict | None = None,
) -> Path:
    """Return the correct artefact folder for this company+role.

    Logic:
    1. If the company root has JD.md directly (legacy single-role) and either
       no role_title is given or role_title matches the tracker record → return root.
    2. If role_title is given → return ARTEFACTS_DIR/Company/role-slug/.
    3. If multiple tracker records exist for this company and role_title is None → raise.
    4. Otherwise return the company root.
    """
    company_root = ARTEFACTS_DIR / company

    # Check how many tracker records exist for this company
    if tracker is None:
        tracker = _load_tracker()
    company_apps = [
        a for a in tracker.get("applications", [])
        if a["company"].lower() == company.lower()
    ]

    # Legacy root: JD.md exists at root with no sub-folders
    root_jd = company_root / "JD.md"
    has_root_jd = root_jd.exists()

    if role_title is None:
        if len(company_apps) > 1:
            raise AmbiguousRoleError(company, [a["role_title"] for a in company_apps])
        # Single or no record — use root
        return company_root

    # role_title given — check if it matches the legacy root record
    if has_root_jd and len(company_apps) <= 1:
        return company_root  # legacy single-role, use root

    # Multi-role: use slug sub-folder
    slug = _make_role_slug(role_title)
    return company_root / slug


def _find_application(tracker: dict, company: str, role_title: str | None) -> dict | None:
    """Find an application record by company (and optionally role_title)."""
    for app in tracker.get("applications", []):
        if app["company"].lower() != company.lower():
            continue
        if role_title is None or app["role_title"].lower() == role_title.lower():
            return app
    return None


def _create_application_record(company: str, role_title: str, jd_path: str) -> dict:
    """Build a new tracker application record in the 'new' stage."""
    now = _utc_now()
    return {
        "id": str(uuid.uuid4()),
        "company": company,
        "role_title": role_title,
        "jd_path": jd_path,
        "stage": "new",
        "date_created": now,
        "history": [{"stage": "new", "at": now}],
        "followups": [],
        "outputs": {},
        "submitted": {},
    }


def _record_output(tracker: dict, company: str, role_title: str | None,
                   output_type: str, entry: dict) -> None:
    """Append an output entry to the tracker record for the given application.

    No-op if the application record is not found (legacy folders without tracker records).
    """
    app = _find_application(tracker, company, role_title)
    if app is not None:
        app.setdefault("outputs", {}).setdefault(output_type, []).append(entry)


# ---------------------------------------------------------------------------
# Follow-up helpers
# ---------------------------------------------------------------------------

def _days_from_now_utc(days: int) -> str:
    """Return date N days from today in YYYY-MM-DD format (UTC)."""
    from datetime import timedelta
    return (datetime.now(timezone.utc) + timedelta(days=days)).strftime("%Y-%m-%d")


def _auto_create_followup(app: dict, new_stage: str) -> None:
    """Create follow-up records on stage transition (Req 11)."""
    followups = app.setdefault("followups", [])

    if new_stage == "applied":
        # send_follow_up_email — deduplicate
        already = any(
            f["action_type"] == "send_follow_up_email" and f["status"] == "pending"
            for f in followups
        )
        if not already:
            followups.append({
                "id": str(uuid.uuid4()),
                "action_type": "send_follow_up_email",
                "due_date": _days_from_now_utc(7),
                "status": "pending",
                "completed_at": None,
            })

    elif new_stage in INTERVIEW_STAGES:
        # send_thank_you_note — deduplicate
        already = any(
            f["action_type"] == "send_thank_you_note" and f["status"] == "pending"
            for f in followups
        )
        if not already:
            followups.append({
                "id": str(uuid.uuid4()),
                "action_type": "send_thank_you_note",
                "due_date": _days_from_now_utc(1),
                "status": "pending",
                "completed_at": None,
            })
        # Cancel any pending follow-up email (Req 11.7)
        _cancel_followup_emails(app)


def _cancel_followup_emails(app: dict) -> None:
    """Set pending send_follow_up_email records to cancelled."""
    for f in app.get("followups", []):
        if f["action_type"] == "send_follow_up_email" and f["status"] == "pending":
            f["status"] = "cancelled"


# ---------------------------------------------------------------------------
# MCP server instance
# ---------------------------------------------------------------------------

MCP_HTTP_HOST = "0.0.0.0"
MCP_HTTP_PORT = 8086
MCP_PUBLIC_HOST = os.environ.get("MCP_PUBLIC_HOST", "localhost")  # externally reachable host:port for OAuth resource metadata


class _StaticBearerVerifier(TokenVerifier):
    """Validates the single shared MCP_AUTH_TOKEN as a static bearer token."""

    async def verify_token(self, token: str) -> AccessToken | None:
        if MCP_AUTH_TOKEN and token == MCP_AUTH_TOKEN:
            return AccessToken(token=token, client_id="job-applications-client", scopes=[])
        return None


if MCP_MODE == "http":
    if not MCP_AUTH_TOKEN:
        sys.stderr.write(
            "[job-applications] ERROR: MCP_MODE=http requires MCP_AUTH_TOKEN to be set.\n"
        )
        sys.exit(1)
    _resource_url = f"http://{MCP_PUBLIC_HOST}:{MCP_HTTP_PORT}"
    mcp = FastMCP(
        "job-applications",
        host=MCP_HTTP_HOST,
        port=MCP_HTTP_PORT,
        token_verifier=_StaticBearerVerifier(),
        auth=AuthSettings(issuer_url=_resource_url, resource_server_url=_resource_url),
    )
else:
    mcp = FastMCP("job-applications")


def _company_dir(company: str) -> Path:
    """Return the path to a company folder, creating it if needed."""
    d = ARTEFACTS_DIR / company
    d.mkdir(parents=True, exist_ok=True)
    return d


def _extract_pdf_text(pdf_path: Path) -> str:
    """Extract text from a PDF file."""
    if PdfReader is None:
        return "[ERROR: PyPDF2 not installed. Install with: pip install PyPDF2]"
    reader = PdfReader(str(pdf_path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _read_file(path: Path) -> str | None:
    """Read a text file, returning None if it doesn't exist."""
    if not path.exists():
        return None
    return path.read_text(encoding="utf-8")


def _find_cv(company_dir: Path) -> Path | None:
    """Find the CV file in a company folder (PDF or Markdown)."""
    for pattern in ["CV*.pdf", "CV*.md", "cv*.pdf", "cv*.md", "Resume*.pdf", "Resume*.md"]:
        matches = list(company_dir.glob(pattern))
        if matches:
            return matches[0]
    # Also check for exact-name CV files
    for f in company_dir.iterdir():
        if f.is_file() and "cv" in f.name.lower() and f.suffix in (".pdf", ".md"):
            return f
    return None


# ---------------------------------------------------------------------------
# JD ingestion (Task 4)
# ---------------------------------------------------------------------------

def _ingest_jd_url(url: str) -> str:
    """Fetch a job description page and return its visible text.

    Raises requests.RequestException (timeout, connection error, or non-200
    status via raise_for_status) — callers are expected to catch this.
    """
    resp = requests.get(url, timeout=15)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    return soup.get_text(separator="\n", strip=True)


def _parse_jd_fields(raw_text: str) -> dict:
    """Regex heuristics to pull structured fields out of raw JD text.

    Fields that can't be confidently detected are returned as None (or an
    empty list for the two skill fields) rather than guessed.
    """
    text = raw_text or ""

    def _search(patterns, flags=re.IGNORECASE):
        for pattern in patterns:
            m = re.search(pattern, text, flags)
            if m:
                return m.group(1).strip()
        return None

    def _extract_list_section(headers):
        pattern = r"(?:%s)\s*[:\-]?\s*\n((?:[ \t]*[-*•][ \t]*.+\n?)+)" % "|".join(headers)
        m = re.search(pattern, text, re.IGNORECASE)
        if not m:
            return []
        lines = m.group(1).strip().split("\n")
        return [re.sub(r"^[ \t]*[-*•][ \t]*", "", line).strip() for line in lines if line.strip()]

    role_title = _search([r"(?:job title|position|role)\s*[:\-]\s*(.+)"])
    if not role_title:
        first_line = text.strip().split("\n")[0] if text.strip() else ""
        if first_line and len(first_line) < 200:
            role_title = first_line.strip("# ").strip()

    company_name = _search([r"(?:company|employer|organi[sz]ation)\s*[:\-]\s*(.+)"])
    location = _search([r"location\s*[:\-]\s*(.+)"])

    employment_type = _search([r"(?:employment type|job type)\s*[:\-]\s*(.+)"])
    if not employment_type:
        m = re.search(r"\b(full[\s-]?time|part[\s-]?time|contract|internship|temporary)\b", text, re.IGNORECASE)
        employment_type = m.group(1) if m else None

    years_match = re.search(r"(\d+)\+?\s*(?:to\s*\d+\s*)?years?\s*(?:of\s*)?experience", text, re.IGNORECASE)
    years_of_experience = int(years_match.group(1)) if years_match else None

    required_skills = _extract_list_section(["required skills", "requirements", "must have"])
    preferred_skills = _extract_list_section(["preferred skills", "nice to have", "bonus"])

    responsibilities_summary = _search([
        r"(?:responsibilities|what you.ll do|role summary)\s*[:\-]?\s*\n(.+?)(?:\n\s*\n|\Z)",
    ], flags=re.IGNORECASE | re.DOTALL)
    if responsibilities_summary:
        responsibilities_summary = " ".join(responsibilities_summary.split())[:1000]

    return {
        "role_title": role_title,
        "company_name": company_name,
        "location": location,
        "employment_type": employment_type,
        "required_skills": required_skills,
        "preferred_skills": preferred_skills,
        "years_of_experience": years_of_experience,
        "responsibilities_summary": responsibilities_summary,
    }


# ---------------------------------------------------------------------------
# Profile manager (Task 5)
# ---------------------------------------------------------------------------

_MONTH_ABBR = {
    "jan": "01", "feb": "02", "mar": "03", "apr": "04", "may": "05", "jun": "06",
    "jul": "07", "aug": "08", "sep": "09", "oct": "10", "nov": "11", "dec": "12",
}


def _normalise_profile_date(token: str) -> str | None:
    """Convert a date token ('Sep 2024', '2019', 'Present') to 'YYYY-MM'/'YYYY'/'present'."""
    token = (token or "").strip()
    if not token:
        return None
    if token.lower() == "present":
        return "present"
    m = re.match(r"^([A-Za-z]{3,})\.?\s+(\d{4})$", token)
    if m:
        month = _MONTH_ABBR.get(m.group(1)[:3].lower())
        if month:
            return f"{m.group(2)}-{month}"
    m = re.match(r"^(\d{4})-(\d{2})$", token)
    if m:
        return token
    m = re.match(r"^(\d{4})$", token)
    if m:
        return m.group(1)
    return None


def _parse_ym(token: str) -> tuple | None:
    """Parse a 'YYYY-MM', 'YYYY', or 'present' token into a (year, month) tuple."""
    token = (token or "").strip()
    if not token:
        return None
    if token.lower() == "present":
        now = datetime.now(timezone.utc)
        return (now.year, now.month)
    m = re.match(r"^(\d{4})-(\d{2})$", token)
    if m:
        return (int(m.group(1)), int(m.group(2)))
    m = re.match(r"^(\d{4})$", token)
    if m:
        return (int(m.group(1)), 1)
    return None


def _compute_years_of_experience(work_exp: list) -> int:
    """Sum non-overlapping work_experience durations in months, floored to whole years."""
    intervals = []
    for entry in work_exp:
        start = _parse_ym(entry.get("start"))
        end = _parse_ym(entry.get("end"))
        if not start or not end:
            continue
        start_idx, end_idx = start[0] * 12 + start[1], end[0] * 12 + end[1]
        if end_idx >= start_idx:
            intervals.append((start_idx, end_idx))

    if not intervals:
        return 0

    intervals.sort()
    merged = [intervals[0]]
    for s, e in intervals[1:]:
        last_s, last_e = merged[-1]
        if s <= last_e:
            merged[-1] = (last_s, max(last_e, e))
        else:
            merged.append((s, e))

    return sum(e - s for s, e in merged) // 12


def _top_n_skills(work_exp: list, skills: list, n: int = 5) -> list:
    """Rank skills by frequency across the explicit skills list and work_experience descriptions."""
    counts = Counter()
    order = []
    for skill in skills:
        key = skill.strip().lower()
        if not key:
            continue
        if key not in counts:
            order.append(key)
        counts[key] += 1
    for entry in work_exp:
        desc = (entry.get("description") or "").lower()
        for key in order:
            if key in desc:
                counts[key] += 1
    ranked = sorted(order, key=lambda k: counts[k], reverse=True)
    return ranked[:n]


def _merge_profile_section(existing: list, incoming: list, source: str, key_fields: tuple) -> tuple:
    """Merge incoming list-section entries into existing by key_fields match.

    LinkedIn-sourced incoming values overwrite conflicting existing fields and
    are logged; cv/session sources only add new entries or fill gaps on
    fields the existing entry left empty. Returns (merged_list, conflicts, changed).
    """
    merged = [dict(e) for e in existing]
    conflicts = []
    changed = False

    def _match(entry, item):
        return all(
            str(entry.get(f, "")).strip().lower() == str(item.get(f, "")).strip().lower()
            for f in key_fields
        )

    for item in incoming:
        target = next((e for e in merged if _match(e, item)), None)
        if target is None:
            new_entry = dict(item)
            new_entry["_source"] = source
            merged.append(new_entry)
            changed = True
            continue

        key_label = ".".join(str(item.get(f, "")) for f in key_fields)
        for field, value in item.items():
            if field in key_fields or field == "_source" or value in (None, "", []):
                continue
            existing_value = target.get(field)
            if existing_value == value:
                continue
            if source == "linkedin":
                if existing_value not in (None, "", []):
                    conflicts.append({
                        "field_path": f"{key_label}.{field}",
                        "linkedin_value": value,
                        "cv_value": existing_value,
                        "flagged_at": _utc_now(),
                    })
                target[field] = value
                target["_source"] = "linkedin"
                changed = True
            elif existing_value in (None, "", []):
                target[field] = value
                changed = True

    return merged, conflicts, changed


def _merge_profile(existing: dict, incoming: dict, source: str) -> dict:
    """Merge an incoming profile fragment into the existing profile.json dict.

    LinkedIn always wins on conflicting list-section fields; the superseded
    value is logged to conflicts[]. Skills are unioned, deduped, and lowercased
    regardless of source.
    """
    profile = dict(existing)
    profile.setdefault("schema_version", "1.0")
    profile["conflicts"] = list(profile.get("conflicts", []))
    profile["last_updated"] = dict(profile.get("last_updated", {}))

    if incoming.get("headline"):
        profile["headline"] = incoming["headline"]
    if incoming.get("current_role"):
        profile["current_role"] = incoming["current_role"]

    for section, key_fields in (
        ("work_experience", ("company", "title")),
        ("education", ("institution", "degree")),
        ("certifications", ("name",)),
    ):
        merged, conflicts, changed = _merge_profile_section(
            profile.get(section, []), incoming.get(section, []), source, key_fields,
        )
        profile[section] = merged
        profile["conflicts"].extend(conflicts)
        if changed:
            profile["last_updated"][section] = _utc_now()

    incoming_skills = incoming.get("skills") or []
    if incoming_skills:
        existing_skills = set(profile.get("skills", []))
        new_skills = {s.strip().lower() for s in incoming_skills if s.strip()}
        if not new_skills <= existing_skills:
            profile["last_updated"]["skills"] = _utc_now()
        profile["skills"] = sorted(existing_skills | new_skills)

    return profile


def _seed_profile_from_cv(cv_path: Path) -> dict:
    """Parse this project's Markdown CV format into profile fields.

    Expects '### Company | Title | Start – End' experience headers, an
    '## EDUCATION...' section of '- **Degree** — Institution' bullets, and a
    '## TECHNICAL SKILLS' section of '- **Category:** skill, skill' bullets.
    """
    text = cv_path.read_text(encoding="utf-8")

    headline = None
    m = re.search(r"^##\s+(.+)$", text, re.MULTILINE)
    if m:
        headline = m.group(1).strip()

    work_experience = []
    exp_pattern = re.compile(
        r"^###\s+(.+?)\s*\|\s*(.+?)\s*\|\s*(.+?)$\n((?:(?!^###|^##).*\n?)*)",
        re.MULTILINE,
    )
    for match in exp_pattern.finditer(text):
        company, title, date_range, body = match.groups()
        parts = re.split(r"\s*[–-]\s*", date_range.strip(), maxsplit=1)
        start = _normalise_profile_date(parts[0]) if parts else None
        end = _normalise_profile_date(parts[1]) if len(parts) > 1 else None
        description = " ".join(
            line.strip(" -*") for line in body.strip().split("\n") if line.strip()
        )
        work_experience.append({
            "title": title.strip(),
            "company": company.strip(),
            "start": start,
            "end": end,
            "description": description,
        })

    current_role = None
    for entry in work_experience:
        if (entry.get("end") or "").lower() == "present":
            current_role = {"title": entry["title"], "company": entry["company"]}
            break

    education = []
    edu_section = re.search(
        r"##\s*EDUCATION.*?\n(.+?)(?:\n---|\n##|\Z)", text, re.DOTALL | re.IGNORECASE,
    )
    if edu_section:
        for line in edu_section.group(1).split("\n"):
            edu_match = re.match(r"^\s*[-*]\s*\*\*(.+?)\*\*\s*[—-]\s*(.+)", line.strip())
            if not edu_match:
                continue
            degree_raw, institution = edu_match.groups()
            field_match = re.search(r"\((.+?)\)", degree_raw)
            degree = re.sub(r"\s*\(.+?\)", "", degree_raw).strip()
            education.append({
                "institution": institution.strip(),
                "degree": degree,
                "field": field_match.group(1).strip() if field_match else None,
                "start": None,
                "end": None,
            })

    skills = []
    skills_section = re.search(
        r"##\s*TECHNICAL SKILLS\s*\n(.+?)(?:\n---|\n##|\Z)", text, re.DOTALL | re.IGNORECASE,
    )
    if skills_section:
        for line in skills_section.group(1).split("\n"):
            skill_match = re.match(r"^\s*[-*]\s*\*\*.+?:\*\*\s*(.+)", line.strip())
            if not skill_match:
                continue
            for item in skill_match.group(1).split(","):
                cleaned = re.sub(r"\(.+?\)", "", item).strip().lower()
                if cleaned:
                    skills.append(cleaned)

    return {
        "headline": headline,
        "current_role": current_role,
        "work_experience": work_experience,
        "education": education,
        "certifications": [],
        "skills": skills,
    }


def _parse_linkedin_export(text: str) -> dict:
    """Heuristic parse of a pasted/exported LinkedIn profile into profile fields.

    Expects loosely-structured plain text with 'Experience' / 'Education' /
    'Skills' section headers on their own line, each followed by blank-line-
    separated entries — the shape produced by copy-pasting a LinkedIn profile page.
    """
    text = text or ""
    lines = text.split("\n")

    section_starts = {}
    for i, line in enumerate(lines):
        stripped = line.strip().lower()
        if stripped in ("experience", "education", "skills"):
            section_starts[stripped] = i

    def _section_body(name):
        if name not in section_starts:
            return []
        start = section_starts[name] + 1
        later = sorted(v for v in section_starts.values() if v > section_starts[name])
        end = later[0] if later else len(lines)
        return lines[start:end]

    def _blocks(body_lines):
        blocks, current = [], []
        for line in body_lines:
            if line.strip():
                current.append(line.strip())
            elif current:
                blocks.append(current)
                current = []
        if current:
            blocks.append(current)
        return blocks

    work_experience = []
    for block in _blocks(_section_body("experience")):
        if len(block) < 2:
            continue
        title, company = block[0], block[1]
        start = end = None
        desc_lines = block[2:]
        if len(block) > 2:
            date_match = re.match(r"^(.+?)\s*[–-]\s*(.+?)(?:\s*\(.*\))?$", block[2])
            if date_match:
                start = _normalise_profile_date(date_match.group(1))
                end = _normalise_profile_date(date_match.group(2))
                desc_lines = block[3:]
        work_experience.append({
            "title": title,
            "company": company,
            "start": start,
            "end": end,
            "description": " ".join(desc_lines),
        })

    education = []
    for block in _blocks(_section_body("education")):
        if not block:
            continue
        institution = block[0]
        degree = field = start = end = None
        if len(block) > 1:
            deg_parts = block[1].split(",", 1)
            degree = deg_parts[0].strip()
            field = deg_parts[1].strip() if len(deg_parts) > 1 else None
        if len(block) > 2:
            date_match = re.match(r"^(\d{4})\s*[–-]\s*(\d{4}|present)$", block[2], re.IGNORECASE)
            if date_match:
                start, end = date_match.group(1), date_match.group(2).lower()
        education.append({
            "institution": institution,
            "degree": degree,
            "field": field,
            "start": start,
            "end": end,
        })

    skills = []
    for line in _section_body("skills"):
        if line.strip():
            skills.extend(s.strip().lower() for s in line.split(",") if s.strip())

    non_blank = [l.strip() for l in lines if l.strip()]
    headline = None
    if len(non_blank) > 1 and non_blank[1].lower() not in ("experience", "education", "skills"):
        headline = non_blank[1]

    current_role = None
    for entry in work_experience:
        if (entry.get("end") or "").lower() == "present":
            current_role = {"title": entry["title"], "company": entry["company"]}
            break

    return {
        "headline": headline,
        "current_role": current_role,
        "work_experience": work_experience,
        "education": education,
        "certifications": [],
        "skills": skills,
    }


# ---------------------------------------------------------------------------
# Templates
# ---------------------------------------------------------------------------

RESEARCH_TEMPLATE = """# {company} — Company Research

> Generated: {date} | Focus: {focus}

## 1. Company Overview
- Full name, ticker, headquarters
- Founded year, CEO, leadership team
- Core business model and revenue streams
- Employee count and global footprint

## 2. Financial Performance
- Annual revenue and growth trajectory
- Key financial metrics (ARR, margins, etc.)
- Recent earnings call highlights

## 3. Strategy & Positioning
- Stated strategic priorities (from annual report / 10-K)
- Key growth initiatives and bets
- Market positioning vs. competitors
- M&A activity and partnerships

## 4. Singapore / APAC Presence
- Singapore office: location, headcount, key leaders
- APAC revenue contribution
- Key accounts in the region
- Government and public sector relationships

## 5. Management Team
- CEO, CFO, CTO, CRO and other C-suite
- Regional/APAC leadership
- Board of directors (notable names)

## 6. Competitors
- Direct competitors and market share
- Competitive advantages and moats
- Recent competitive moves

## 7. Employee & Partner Sentiment
- Glassdoor/Indeed ratings and themes
- Common employee complaints and praise
- Partner/channel feedback
- Recent layoffs or hiring freezes

## 8. Procurement Footprint (Singapore Government)
- Government contracts and tenders
- Key agency relationships

## 9. Key Insights for Interview
- What problems is this company trying to solve?
- Where does my experience align?
- What questions should I ask?

---
*Research focus: {focus}*
"""

TERRITORY_TEMPLATE = """# {company} — Territory & Contact Map

> Generated: {date}

## Account: {account}

| Name | Role | Email | Phone | Organization | Relationship | Source |
|------|------|-------|-------|--------------|---------------|--------|
| (contacts will be listed here) | | | | | | |

### Warm Paths
- (Direct relationships that can provide introductions)

### Account Intelligence
- (Key decisions, procurement history, strategic priorities)

---
"""

COVER_LETTER_TONES = {
    "bold": "Bold and direct. Lead with your strongest differentiator. Use confident, assertive language. Take a clear position on why you're the right candidate.",
    "conservative": "Professional and measured. Follow traditional cover letter conventions. Emphasize qualifications and experience systematically.",
    "storyteller": "Narrative-driven. Weave your career story into a compelling arc that connects to the role. Use the 'Ground Truth' framing if relevant.",
}

PITCH_FORMATS = {
    "narrative": "Flowing narrative with a clear arc: opening hook, key messages, closing ask. Conversational tone.",
    "bullet_points": "Structured bullet points organized by topic. Quick to scan, easy to reference during interview.",
    "star_stories": "Situation-Task-Action-Result format for each key experience. Proves claims with evidence.",
}

# Requirement 6.5: numeric/keyword segments from the Base_CV that must be
# copied verbatim into any Tailored_CV — never altered or fabricated.
PROTECTED_NUMERIC_RE = re.compile(r"\d+\s*(?:%|\$|SGD|€|£)", re.IGNORECASE)
PROTECTED_KEYWORD_RE = re.compile(r"\b(?:ARR|quota|deal|target)\b", re.IGNORECASE)

VALID_DIFF_CHANGE_TYPES = {"reorder", "condense", "add", "remove", "replace"}


def _protected_lines(base_cv_text: str) -> list[str]:
    """Base_CV lines containing a quantified figure or protected keyword."""
    lines = []
    for line in base_cv_text.splitlines():
        stripped = line.strip()
        if stripped and (PROTECTED_NUMERIC_RE.search(stripped) or PROTECTED_KEYWORD_RE.search(stripped)):
            lines.append(stripped)
    return lines


def _resolve_base_cv_content(company_dir: Path) -> str | None:
    """Global Base_CV first (BASE_CV_PATH); falls back to a per-company CV file."""
    if BASE_CV_PATH.exists():
        return _read_file(BASE_CV_PATH)
    cv_file = _find_cv(company_dir)
    if cv_file is None:
        return None
    return _extract_pdf_text(cv_file) if cv_file.suffix.lower() == ".pdf" else _read_file(cv_file)


def _generate_diff_summary(base_text: str, new_text: str) -> list[dict]:
    """Fallback structural diff (Req 6.8) used when the caller supplies none."""
    base_lines = base_text.splitlines()
    new_lines = new_text.splitlines()
    matcher = difflib.SequenceMatcher(a=base_lines, b=new_lines, autojunk=False)
    entries = []
    current_section = "General"
    scanned = 0
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        for line in base_lines[scanned:i1]:
            if line.strip().startswith("#"):
                current_section = line.strip("# ").strip()
        scanned = i2 if tag != "equal" else i2
        if tag == "equal":
            continue
        change_type = {"replace": "replace", "delete": "remove", "insert": "add"}[tag]
        added = " / ".join(l.strip() for l in new_lines[j1:j2] if l.strip())
        removed = " / ".join(l.strip() for l in base_lines[i1:i2] if l.strip())
        description = (added or removed or "content changed")[:200]
        entries.append({"section": current_section, "change_type": change_type, "description": description})
    return entries


# ---------------------------------------------------------------------------
# Export Engine (Task 8)
# ---------------------------------------------------------------------------

DOCUMENT_TYPE_FILES = {"tailored_cv": "CV_tailored.md", "cover_letter": "Cover_Letter.md"}
EXPORT_FORMATS = {"pdf", "docx"}

EXPORT_DEP_INSTALL = {
    "markdown": "pip install markdown~=3.6",
    "weasyprint": "pip install weasyprint~=61.0",
    "python-docx": "pip install python-docx~=1.1",
}

PDF_CSS = """
@page { size: A4; margin: 2.5cm; }
body { font-family: Georgia, serif; font-size: 11pt; line-height: 1.4; }
h1, h2, h3 { font-weight: bold; }
h1 { font-size: 20pt; }
h2 { font-size: 15pt; }
h3 { font-size: 12pt; }
"""


def _check_export_deps(fmt: str) -> list[str]:
    """Return pip install commands for any dependency missing for `fmt`."""
    missing = []
    if fmt == "pdf":
        try:
            import markdown  # noqa: F401
        except ImportError:
            missing.append(EXPORT_DEP_INSTALL["markdown"])
        try:
            import weasyprint  # noqa: F401
        except Exception:
            missing.append(EXPORT_DEP_INSTALL["weasyprint"])
    elif fmt == "docx":
        try:
            import docx  # noqa: F401
        except ImportError:
            missing.append(EXPORT_DEP_INSTALL["python-docx"])
    return missing


def _export_to_pdf(source_md: Path, output_path: Path) -> None:
    """Render a Markdown file to PDF (Req 8.3): A4, 2.5cm margins, Georgia 11pt."""
    import markdown
    import weasyprint

    md_text = source_md.read_text(encoding="utf-8")
    body_html = markdown.markdown(md_text, extensions=["tables", "nl2br"])
    html_doc = f"<html><head><style>{PDF_CSS}</style></head><body>{body_html}</body></html>"
    weasyprint.HTML(string=html_doc).write_pdf(str(output_path))


def _export_to_docx(source_md: Path, output_path: Path) -> None:
    """Render a Markdown file to DOCX (Req 8.4): 2.54cm margins, Calibri body."""
    from docx import Document
    from docx.shared import Cm, Pt

    document = Document()
    for section in document.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    heading_sizes = {1: (16, True), 2: (13, True), 3: (11, True)}
    md_text = source_md.read_text(encoding="utf-8")
    for line in md_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        heading_match = re.match(r"^(#{1,3})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            size, bold = heading_sizes[level]
            para = document.add_paragraph()
            run = para.add_run(heading_match.group(2))
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.bold = bold
            if level == 3:
                run.font.italic = True
        elif stripped.startswith(("- ", "* ")):
            para = document.add_paragraph(stripped[2:], style="List Bullet")
            para.paragraph_format.left_indent = Cm(0.5)
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)
        else:
            para = document.add_paragraph(stripped)
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)

    document.save(str(output_path))


# ---------------------------------------------------------------------------
# MCP Tools
# ---------------------------------------------------------------------------


@mcp.tool()
def create_application(company: str, jd_path: str, role_title: str | None = None) -> dict:
    """Create a new job application folder and parse the job description.

    Creates the company folder, extracts text from the JD (PDF or Markdown),
    and writes JD.md. Returns structured data for setting up CRM records.

    After calling this tool, use the ai-crm create_account and create_opportunity
    tools to set up CRM tracking.

    Args:
        company: Target employer name (e.g., "Gartner"). Used as the folder name.
        jd_path: Path to the job description file (PDF or Markdown).
        role_title: Optional role title override. If omitted, extracted from JD.
    """
    tracker = _load_tracker()
    try:
        # Resolve using the caller-supplied role_title only — an extracted-from-JD
        # fallback must not influence folder placement (Req 10, legacy compat).
        company_dir = _resolve_company_folder(company, role_title, tracker)
    except AmbiguousRoleError as e:
        return {"ok": False, "error": "ambiguous_role", "company": e.company, "roles": e.roles}
    company_dir.mkdir(parents=True, exist_ok=True)

    jd_source = Path(jd_path)

    # Read and parse the JD
    if not jd_source.exists():
        return {"error": f"JD file not found: {jd_path}"}

    if jd_source.suffix.lower() == ".pdf":
        jd_text = _extract_pdf_text(jd_source)
    elif jd_source.suffix.lower() in (".md", ".markdown", ".txt"):
        jd_text = jd_source.read_text(encoding="utf-8")
    else:
        return {"error": f"Unsupported JD format: {jd_source.suffix}. Use PDF or Markdown."}

    # Extract role title from first line or filename if not provided (display only)
    if not role_title:
        first_line = jd_text.strip().split("\n")[0] if jd_text else ""
        if first_line and len(first_line) < 200:
            role_title = first_line.strip("# ").strip()
        else:
            role_title = jd_source.stem.replace("_", " ").replace("-", " ")

    # Write JD.md
    jd_md = company_dir / "JD.md"
    jd_md.write_text(jd_text, encoding="utf-8")

    # Check for existing CV
    cv_file = _find_cv(company_dir)

    return {
        "company": company,
        "folder_path": str(company_dir),
        "role_title": role_title,
        "jd_path": str(jd_md),
        "jd_length": len(jd_text),
        "cv_found": cv_file is not None,
        "cv_path": str(cv_file) if cv_file else None,
        "existing_files": [f.name for f in company_dir.iterdir() if f.is_file()],
        "next_steps": [
            "1. Use the deep-research skill to research the company",
            "2. Call company_research to save research results",
            "3. Use ai-crm create_account to create the company account",
            "4. Use ai-crm create_opportunity to create the job opportunity",
        ],
    }


@mcp.tool()
def get_application_status(company: str, role_title: str | None = None) -> dict:
    """Check the status of a job application workflow.

    Returns which files exist, what's been completed, and what steps remain.

    Args:
        company: Target employer name (e.g., "Gartner").
        role_title: Optional role title, required to disambiguate companies
            with multiple tracked roles.
    """
    tracker = _load_tracker()
    try:
        company_dir = _resolve_company_folder(company, role_title, tracker)
    except AmbiguousRoleError as e:
        return {"ok": False, "error": "ambiguous_role", "company": e.company, "roles": e.roles}

    app = _find_application(tracker, company, role_title)

    if not company_dir.exists():
        return {
            "company": company,
            "exists": False,
            "next_steps": [
                "Call create_application first to set up the company folder.",
            ],
        }

    files = {f.name: f for f in company_dir.iterdir() if f.is_file()}
    cv_file = _find_cv(company_dir)

    # Check which workflow steps are completed
    steps = {
        "job_description": "JD.md" in files,
        "research": "research.md" in files,
        "territory_map": "territory_map.md" in files,
        "cover_letter": any("cover" in f.lower() or "letter" in f.lower() for f in files),
        "pitch": "pitch.md" in files,
        "tailored_cv": "CV_tailored.md" in files,
        "gap_analysis": "gap_analysis.md" in files,
        "learning_program": "learning_program.md" in files,
        "interview_notes": "interview_notes.md" in files,
    }

    completed = [step for step, done in steps.items() if done]
    pending = [step for step, done in steps.items() if not done]

    # Determine next recommended action
    if not steps["job_description"]:
        next_action = "create_application"
    elif not steps["research"]:
        next_action = "company_research"
    elif not steps["territory_map"]:
        next_action = "map_territory"
    elif not steps["pitch"]:
        next_action = "generate_pitch"
    elif not steps["cover_letter"]:
        next_action = "generate_cover_letter"
    elif not steps["tailored_cv"]:
        next_action = "tailor_cv"
    elif not steps["gap_analysis"]:
        next_action = "analyse_gaps"
    elif not steps["learning_program"]:
        next_action = "generate_learning_program"
    else:
        next_action = "All steps completed!"

    result = {
        "company": company,
        "folder_path": str(company_dir),
        "exists": True,
        "steps": steps,
        "completed": completed,
        "pending": pending,
        "next_action": next_action,
        "files": list(files.keys()),
        "cv_available": cv_file is not None,
    }

    # Tracker read-through (Task 3): most companies have no tracker record
    # yet (ingest_jd, which creates them, is Task 4), so this falls back to
    # the filesystem-only view above when there's no match.
    if app is not None:
        result["tracked"] = True
        result["application_id"] = app.get("id")
        result["role_title"] = app.get("role_title")
        result["stage"] = app.get("stage")
        result["date_created"] = app.get("date_created")
        result["followups"] = app.get("followups", [])
        result["outputs"] = app.get("outputs", {})
        result["submitted"] = app.get("submitted", {})
    else:
        result["tracked"] = False
        result["outputs"] = {}
        result["submitted"] = {}

    # Check for submitted/ folder on disk
    submitted_dir = company_dir / "submitted"
    if submitted_dir.is_dir():
        result["submitted_files"] = [f.name for f in submitted_dir.iterdir() if f.is_file()]
    else:
        result["submitted_files"] = []

    return result


@mcp.tool()
def update_stage(company: str, role_title: str, new_stage: str) -> dict:
    """Advance a tracked application to a new pipeline stage.

    Validates the transition against the stage machine (e.g. you can't skip
    from 'new' straight to 'offer', and terminal stages like 'rejected'
    accept no further transitions), appends a history entry, and
    auto-creates/cancels follow-up reminders as appropriate.

    Args:
        company: Target employer name.
        role_title: Role title identifying which tracked application to update.
        new_stage: Target stage — one of: new, applied, screening,
            interview_r1, interview_r2, interview_r3, offer, accepted,
            rejected, withdrawn.
    """
    if new_stage not in VALID_STAGES:
        return {"ok": False, "error": "invalid_stage", "valid_stages": sorted(VALID_STAGES)}

    tracker = _load_tracker()
    app = _find_application(tracker, company, role_title)
    if app is None:
        return {"ok": False, "error": "application_not_found", "company": company, "role_title": role_title}

    current_stage = app["stage"]
    allowed_next = VALID_TRANSITIONS.get(current_stage, set())
    if new_stage not in allowed_next:
        return {
            "ok": False,
            "error": "invalid_transition",
            "current_stage": current_stage,
            "requested_stage": new_stage,
            "valid_next_stages": sorted(allowed_next),
        }

    app["stage"] = new_stage
    app.setdefault("history", []).append({"stage": new_stage, "at": _utc_now()})
    _auto_create_followup(app, new_stage)
    _save_tracker(tracker)

    return {
        "ok": True,
        "company": company,
        "role_title": role_title,
        "previous_stage": current_stage,
        "new_stage": new_stage,
    }


@mcp.tool()
def list_applications(company: str | None = None, stage: str | None = None) -> dict:
    """List tracked job applications, optionally filtered by company and/or stage.

    Results are sorted by date_created descending (most recent first).

    Args:
        company: Optional company name filter (case-insensitive).
        stage: Optional stage filter.
    """
    tracker = _load_tracker()
    apps = tracker.get("applications", [])

    if company is not None:
        apps = [a for a in apps if a["company"].lower() == company.lower()]
    if stage is not None:
        apps = [a for a in apps if a["stage"] == stage]

    apps = sorted(apps, key=lambda a: a["date_created"], reverse=True)

    result = {
        "count": len(apps),
        "applications": [
            {
                "id": a["id"],
                "company": a["company"],
                "role_title": a["role_title"],
                "stage": a["stage"],
                "date_created": a["date_created"],
            }
            for a in apps
        ],
    }

    if company is not None and len({a["role_title"] for a in apps}) > 1:
        result["note"] = (
            f"Multiple roles tracked at {company}; pass role_title to "
            "company-scoped tools to disambiguate."
        )

    return result


@mcp.tool()
def get_due_followups() -> dict:
    """List pending follow-up actions due today or earlier (UTC), sorted ascending.

    Does not mutate stored follow-up status — items past their due date are
    flagged with overdue: true in the response only.
    """
    tracker = _load_tracker()
    today = _days_from_now_utc(0)

    due = []
    for app in tracker.get("applications", []):
        for f in app.get("followups", []):
            if f["status"] != "pending":
                continue
            if f["due_date"] > today:
                continue
            due.append({
                "application_id": app["id"],
                "company": app["company"],
                "role_title": app["role_title"],
                "followup_id": f["id"],
                "action_type": f["action_type"],
                "due_date": f["due_date"],
                "overdue": f["due_date"] < today,
            })

    due.sort(key=lambda d: d["due_date"])
    return {"count": len(due), "due_followups": due}


@mcp.tool()
def mark_followup_complete(application_id: str, followup_id: str) -> dict:
    """Mark a follow-up action as completed.

    Args:
        application_id: The application record's id (from list_applications
            or get_due_followups).
        followup_id: The follow-up entry's id within that application.
    """
    tracker = _load_tracker()
    for app in tracker.get("applications", []):
        if app["id"] != application_id:
            continue
        for f in app.get("followups", []):
            if f["id"] == followup_id:
                f["status"] = "completed"
                f["completed_at"] = _utc_now()
                _save_tracker(tracker)
                return {"ok": True, "application_id": application_id, "followup_id": followup_id}
        return {"ok": False, "error": "followup_not_found", "application_id": application_id, "followup_id": followup_id}
    return {"ok": False, "error": "application_not_found", "application_id": application_id}


@mcp.tool()
def ingest_jd(
    company: str,
    role_title: str,
    jd_path: str | None = None,
    jd_url: str | None = None,
    jd_text: str | None = None,
) -> dict:
    """Ingest a job description and create/update its tracker record.

    Exactly one content source must be given: jd_path (file), jd_url (fetch),
    or jd_text (pasted text). When jd_text is provided, jd_url may also be
    given as a reference/provenance URL (the URL is stored but not fetched).
    Writes JD.md into the resolved company/role folder, creates a new
    'new'-stage tracker record (or updates an existing one, preserving its
    current stage), and returns the regex-extracted structured fields.

    Args:
        company: Target employer name.
        role_title: Role title — used for tracker lookup and folder resolution.
        jd_path: Path to a local JD file (PDF or Markdown). Content source.
        jd_url: URL of a JD posting to fetch and parse. Content source unless
            jd_text is also given, in which case this is stored as a reference.
        jd_text: Pasted JD text. Content source. When given alongside jd_url,
            the URL is recorded as provenance metadata but not fetched.
    """
    content_sources = sum(1 for s in (jd_path, jd_url, jd_text) if s)
    if content_sources == 0:
        return {"ok": False, "error": "no_source_given"}
    if content_sources > 1 and not (jd_text and jd_url and not jd_path):
        # Only jd_text + jd_url is allowed together (pasted text + reference URL).
        # Any other combination of multiple sources is ambiguous.
        if jd_path and jd_url:
            return {"ok": False, "error": "both_sources_given"}
        if jd_path and jd_text:
            return {"ok": False, "error": "both_sources_given",
                    "message": "Provide jd_path or jd_text, not both"}
        return {"ok": False, "error": "both_sources_given"}

    # Determine content source and reference URL
    reference_url = None
    if jd_text:
        # Pasted text — use directly; jd_url (if given) is provenance only
        reference_url = jd_url  # may be None
    elif jd_path:
        jd_source = Path(jd_path)
        if not jd_source.exists():
            return {"ok": False, "error": "file_not_found", "jd_path": jd_path}
        suffix = jd_source.suffix.lower()
        if suffix == ".pdf":
            jd_text = _extract_pdf_text(jd_source)
            if not jd_text.strip():
                return {"ok": False, "error": "zero_text_pdf", "jd_path": jd_path}
        elif suffix in (".md", ".markdown", ".txt"):
            jd_text = jd_source.read_text(encoding="utf-8")
        else:
            return {"ok": False, "error": "unsupported_format", "jd_path": jd_path}
    else:
        # jd_url only — fetch it
        try:
            jd_text = _ingest_jd_url(jd_url)
        except requests.RequestException as e:
            return {"ok": False, "error": "url_error", "detail": str(e), "jd_url": jd_url}

        # Guard against JS-rendered pages that return only a title like "Meta Careers".
        # Only checked for URL-fetched content — users know what they're pasting or
        # providing as a file, but URL scraping can silently return useless content.
        MIN_URL_JD_LENGTH = 50
        if len(jd_text.strip()) < MIN_URL_JD_LENGTH:
            return {
                "ok": False,
                "error": "jd_content_too_short",
                "message": (
                    f"URL returned only {len(jd_text.strip())} characters — likely a page "
                    f"title or navigation text rather than the actual job description. "
                    f"This usually happens when the job posting uses JavaScript rendering "
                    f"(e.g. LinkedIn, Meta, Workday). Try using jd_text to paste the "
                    f"JD content directly instead."
                ),
                "jd_length": len(jd_text.strip()),
                "source": "jd_url",
            }

    tracker = _load_tracker()
    try:
        company_dir = _resolve_company_folder(company, role_title, tracker)
    except AmbiguousRoleError as e:
        return {"ok": False, "error": "ambiguous_role", "company": e.company, "roles": e.roles}
    company_dir.mkdir(parents=True, exist_ok=True)

    jd_md = company_dir / "JD.md"
    jd_md.write_text(jd_text, encoding="utf-8")

    existing = _find_application(tracker, company, role_title)
    if existing is not None:
        existing["jd_path"] = str(jd_md)
        app = existing
    else:
        app = _create_application_record(company, role_title, str(jd_md))
        tracker.setdefault("applications", []).append(app)

    # Store reference URL as provenance metadata when jd_text was used with jd_url
    if reference_url:
        app["jd_source_url"] = reference_url

    _save_tracker(tracker)

    fields = _parse_jd_fields(jd_text)

    result = {
        "ok": True,
        "application_id": app["id"],
        "company": company,
        "role_title": role_title,
        "stage": app["stage"],
        "folder_path": str(company_dir),
        "jd_path": str(jd_md),
        "jd_length": len(jd_text),
        "fields": fields,
    }
    if reference_url:
        result["jd_source_url"] = reference_url
    return result


@mcp.tool()
def update_profile(source: str, cv_path: str | None = None, text: str | None = None) -> dict:
    """Seed or update profile.json from a CV file or freeform session notes.

    Args:
        source: 'cv' to seed/refresh from a Markdown CV, or 'session' to merge
            freeform notes typed during the conversation.
        cv_path: Path to a Markdown CV. Defaults to BASE_CV_PATH when source='cv'.
        text: Freeform profile notes. Required when source='session'.
    """
    if source == "cv":
        path = Path(cv_path) if cv_path else BASE_CV_PATH
        if not path.exists():
            return {"ok": False, "error": "file_not_found", "cv_path": str(path)}
        incoming = _seed_profile_from_cv(path)
    elif source == "session":
        if not text:
            return {"ok": False, "error": "text_required"}
        incoming = _parse_linkedin_export(text)
    else:
        return {"ok": False, "error": "unsupported_source", "source": source}

    existing = _load_profile()
    merged = _merge_profile(existing, incoming, source)
    _save_profile(merged)

    fields_updated = sum(
        1 for section in ("work_experience", "education", "certifications", "skills")
        if json.dumps(existing.get(section), sort_keys=True) != json.dumps(merged.get(section), sort_keys=True)
    )

    return {
        "ok": True,
        "source": source,
        "fields_updated": fields_updated,
        "work_experience_count": len(merged.get("work_experience", [])),
        "education_count": len(merged.get("education", [])),
        "skills_count": len(merged.get("skills", [])),
        "conflicts_count": len(merged.get("conflicts", [])),
    }


@mcp.tool()
def refresh_profile_from_linkedin(
    url: str | None = None,
    file_path: str | None = None,
    text: str | None = None,
) -> dict:
    """Merge a LinkedIn profile export (URL, file, or pasted text) into profile.json.

    Exactly one of url, file_path, or text should be given. LinkedIn-sourced
    values always win over existing CV data on conflict; superseded values are
    logged to profile.json's conflicts[] list.
    """
    sources_given = sum(1 for s in (url, file_path, text) if s)
    if sources_given == 0:
        return {"ok": False, "error": "no_source_given"}
    if sources_given > 1:
        return {"ok": False, "error": "multiple_sources_given"}

    if url:
        try:
            raw_text = _ingest_jd_url(url)
        except requests.RequestException as e:
            return {"ok": False, "error": "url_blocked_or_unreachable", "detail": str(e)}
    elif file_path:
        source_path = Path(file_path)
        if not source_path.exists():
            return {"ok": False, "error": "file_not_found", "file_path": file_path}
        suffix = source_path.suffix.lower()
        if suffix == ".pdf":
            raw_text = _extract_pdf_text(source_path)
        elif suffix in (".txt", ".md"):
            raw_text = source_path.read_text(encoding="utf-8")
        else:
            return {"ok": False, "error": "unsupported_format", "file_path": file_path}
    else:
        raw_text = text

    if not raw_text or not raw_text.strip():
        return {"ok": False, "error": "empty_content"}

    lowered = raw_text.lower()
    if not any(marker in lowered for marker in ("experience", "education", "skills")):
        return {"ok": False, "error": "invalid_linkedin_format"}

    incoming = _parse_linkedin_export(raw_text)
    existing = _load_profile()
    conflicts_before = len(existing.get("conflicts", []))
    merged = _merge_profile(existing, incoming, "linkedin")
    _save_profile(merged)

    return {
        "ok": True,
        "work_experience_extracted": len(incoming.get("work_experience", [])),
        "education_extracted": len(incoming.get("education", [])),
        "skills_extracted": len(incoming.get("skills", [])),
        "conflicts_flagged": len(merged.get("conflicts", [])) - conflicts_before,
    }


@mcp.tool()
def get_profile_summary() -> dict:
    """Return a condensed view of profile.json: current role, experience, top skills, education."""
    profile = _load_profile()
    if not profile.get("work_experience") and not profile.get("skills"):
        return {"ok": True, "setup_required": True}

    return {
        "ok": True,
        "setup_required": False,
        "headline": profile.get("headline"),
        "current_role": profile.get("current_role"),
        "years_of_experience": _compute_years_of_experience(profile.get("work_experience", [])),
        "top_skills": _top_n_skills(profile.get("work_experience", []), profile.get("skills", []), n=5),
        "education": profile.get("education", []),
        "last_updated": profile.get("last_updated", {}),
    }


# ---------------------------------------------------------------------------
# Intelligence context-prep tools (Task 6)
# ---------------------------------------------------------------------------

MATCH_SCORE_WEIGHTS = {
    "required_skills_match": 40,
    "years_of_experience_match": 25,
    "seniority_alignment": 15,
    "industry_domain_alignment": 10,
    "preferred_skills_match": 10,
}

LEARNING_PRIORITY_DAYS = {"high": 30, "medium": 60, "low": 90}


def _load_jd_fields_for_app(app: dict) -> dict | None:
    """Re-parse JD fields from the JD.md referenced by a tracker record.

    Returns None if the record has no jd_path or the file is missing/empty —
    fields are re-derived on demand rather than cached on the tracker record,
    so ingest_jd never needs a companion cache-invalidation path.
    """
    jd_path = app.get("jd_path")
    if not jd_path:
        return None
    jd_text = _read_file(Path(jd_path))
    if not jd_text or not jd_text.strip():
        return None
    return _parse_jd_fields(jd_text)


@mcp.tool()
def score_match(company: str, role_title: str) -> dict:
    """Prepare context for scoring how well the Profile matches a JD.

    Returns JD fields, a profile summary, the five named scoring weights, and
    the expected output schema. Claude performs the actual scoring and calls
    back with save_match_score.

    Args:
        company: Target employer name.
        role_title: Role title identifying the tracked application.
    """
    tracker = _load_tracker()
    app = _find_application(tracker, company, role_title)
    if app is None:
        return {"ok": False, "error": "jd_not_ingested", "company": company, "role_title": role_title}

    jd_fields = _load_jd_fields_for_app(app)
    if jd_fields is None:
        return {"ok": False, "error": "jd_not_ingested", "company": company, "role_title": role_title}

    profile = _load_profile()
    if not profile.get("work_experience") and not profile.get("skills"):
        return {"ok": False, "error": "profile_not_initialised"}

    profile_summary = {
        "headline": profile.get("headline"),
        "current_role": profile.get("current_role"),
        "years_of_experience": _compute_years_of_experience(profile.get("work_experience", [])),
        "skills": profile.get("skills", []),
        "education": profile.get("education", []),
        "work_experience": profile.get("work_experience", []),
    }

    # Use RequirementService to extract requirements from JD
    evidence_service = _get_or_create_evidence_service()
    req_service = RequirementService(evidence_service)
    requirements = req_service.extract_requirements(jd_fields)

    return {
        "ok": True,
        "company": company,
        "role_title": role_title,
        "jd_fields": jd_fields,
        "profile_summary": profile_summary,
        "extracted_requirements": [
            {
                "statement": req.statement,
                "type": req.type.value,
                "confidence": req.confidence.value,
                "source": req.source_jd_field,
            }
            for req in requirements.requirements
        ],
        "weights": MATCH_SCORE_WEIGHTS,
        "instructions": (
            "Score each of the five named dimensions in `weights` from 0-100, then combine "
            "them using those weights to produce an overall Match_Score (0-100). Write a "
            "reasoning section of no more than 500 words, identifying up to 3 strengths and "
            "up to 3 gaps (fewer if fewer exist). Under missing_skills, list every JD "
            "required or preferred skill absent from profile_summary — this feeds "
            "generate_learning_program. Base every judgment only on jd_fields and "
            "profile_summary; do not invent profile content."
        ),
        "output_schema": {
            "overall": "int 0-100",
            "sub_scores": {dim: "int 0-100" for dim in MATCH_SCORE_WEIGHTS},
            "reasoning": "string, <= 500 words",
            "strengths": "list of up to 3 strings",
            "gaps": "list of up to 3 strings",
            "missing_skills": "list of strings",
        },
        "call_next": "save_match_score(company, role_title, overall, sub_scores, reasoning, strengths, gaps, missing_skills)",
    }


@mcp.tool()
def save_match_score(
    company: str,
    role_title: str,
    overall: float,
    sub_scores: dict,
    reasoning: str,
    strengths: list,
    gaps: list,
    missing_skills: list,
) -> dict:
    """Persist a computed Match_Score to the tracker record.

    Args:
        company: Target employer name.
        role_title: Role title identifying the tracked application.
        overall: Combined Match_Score, 0-100 inclusive.
        sub_scores: Dict of the five named dimension scores, each 0-100.
        reasoning: Human-readable explanation (<= 500 words).
        strengths: Up to 3 strength statements.
        gaps: Up to 3 gap statements.
        missing_skills: Required/preferred JD skills absent from the Profile.
    """
    tracker = _load_tracker()
    app = _find_application(tracker, company, role_title)
    if app is None:
        return {"ok": False, "error": "jd_not_ingested", "company": company, "role_title": role_title}

    if not isinstance(overall, (int, float)) or isinstance(overall, bool) or not (0 <= overall <= 100):
        return {"ok": False, "error": "invalid_score", "field": "overall", "value": overall}

    missing_dims = sorted(set(MATCH_SCORE_WEIGHTS) - set(sub_scores or {}))
    if missing_dims:
        return {"ok": False, "error": "invalid_score", "field": "sub_scores", "missing_dimensions": missing_dims}
    for dim in MATCH_SCORE_WEIGHTS:
        val = sub_scores[dim]
        if not isinstance(val, (int, float)) or isinstance(val, bool) or not (0 <= val <= 100):
            return {"ok": False, "error": "invalid_score", "field": f"sub_scores.{dim}", "value": val}

    computed_at = _utc_now()
    app["match_score"] = {
        "overall": overall,
        "sub_scores": sub_scores,
        "reasoning": reasoning,
        "strengths": strengths,
        "gaps": gaps,
        "missing_skills": missing_skills,
        "computed_at": computed_at,
    }
    # Record output in tracker
    _record_output(tracker, company, role_title, "match_score", {
        "overall": overall,
        "saved_at": computed_at,
    })
    _save_tracker(tracker)

    return {
        "ok": True,
        "company": company,
        "role_title": role_title,
        "overall": overall,
        "computed_at": computed_at,
    }


@mcp.tool()
def analyse_gaps(company: str, role_title: str) -> dict:
    """Prepare context for comparing the Base_CV against a JD's requirements.

    Returns the JD fields, Base_CV content, missing_skills from the latest
    Match_Score (if any), and a gap-item schema with an explicit
    no-fabrication instruction. Claude performs the analysis and calls back
    with save_gap_analysis.

    Args:
        company: Target employer name.
        role_title: Role title identifying the tracked application.
    """
    tracker = _load_tracker()
    app = _find_application(tracker, company, role_title)
    if app is None:
        return {"ok": False, "error": "jd_not_ingested", "company": company, "role_title": role_title}

    jd_fields = _load_jd_fields_for_app(app)
    if jd_fields is None:
        return {"ok": False, "error": "jd_not_ingested", "company": company, "role_title": role_title}

    if not BASE_CV_PATH.exists():
        return {"ok": False, "error": "base_cv_not_found", "base_cv_path": str(BASE_CV_PATH)}
    base_cv_content = (
        _extract_pdf_text(BASE_CV_PATH) if BASE_CV_PATH.suffix.lower() == ".pdf"
        else _read_file(BASE_CV_PATH)
    )

    profile = _load_profile()
    if not profile.get("work_experience") and not profile.get("skills"):
        return {"ok": False, "error": "profile_not_initialised"}

    match_score = app.get("match_score")
    missing_skills = match_score.get("missing_skills", []) if match_score else []

    # Use RequirementService to extract and match requirements
    evidence_service = _get_or_create_evidence_service()
    req_service = RequirementService(evidence_service)
    requirements = req_service.extract_requirements(jd_fields)

    # Get requirement matches (empty for now, will be populated in Gate 6)
    evidence_matches = {}
    for req in requirements.requirements:
        evidence_matches[req.requirement_id] = req_service.match_requirement(req)

    # Identify gaps
    gaps = req_service.identify_gaps(requirements, evidence_matches)

    return {
        "ok": True,
        "company": company,
        "role_title": role_title,
        "jd_fields": jd_fields,
        "base_cv_content": base_cv_content,
        "missing_skills": missing_skills,
        "match_score_available": match_score is not None,
        "extracted_gaps": [
            {
                "requirement": gap.requirement_statement,
                "type": gap.type.value,
                "status": gap.status.value,
                "reasoning": gap.reasoning,
            }
            for gap in gaps
        ],
        "gap_schema": {
            "gap_id": "unique string identifier",
            "category": "missing | understated | mismatch",
            "jd_criterion": "string — the JD requirement this gap addresses",
            "affected_cv_section": "string — CV role/section this relates to",
            "current_text_excerpt": "verbatim CV text, or null when category is 'missing'",
            "recommendation": "specific addition/reframing/reordering instruction",
        },
        "instructions": (
            "Compare the JD's required and preferred criteria in jd_fields against "
            "base_cv_content. Classify each gap as 'missing' (absent from the CV entirely), "
            "'understated' (present but not among the first two bullets of the most relevant "
            "role section), or 'mismatch' (present but framed for a different industry or "
            "context than the JD). For 'missing' gaps set current_text_excerpt to null and "
            "give the recommended addition text as the recommendation. For 'understated' or "
            "'mismatch' gaps, current_text_excerpt must be copied verbatim from "
            "base_cv_content, with recommendation as a specific replacement or reordering "
            "instruction. Do NOT fabricate experience, credentials, or achievements — only "
            "suggest reframing, reordering, or expanding on content already present in "
            "base_cv_content or the Profile_Store."
        ),
        "output_path": str(_resolve_company_folder(company, role_title, tracker) / "gap_analysis.md"),
        "call_next": "save_gap_analysis(company, role_title, gaps)",
    }


@mcp.tool()
def save_gap_analysis(company: str, role_title: str, gaps: list) -> dict:
    """Validate and save a Gap_Analysis to {Company_Folder}/gap_analysis.md.

    Args:
        company: Target employer name.
        role_title: Role title identifying the tracked application.
        gaps: List of gap items (gap_id, category, jd_criterion,
            affected_cv_section, current_text_excerpt, recommendation).
    """
    tracker = _load_tracker()
    app = _find_application(tracker, company, role_title)
    if app is None:
        return {"ok": False, "error": "jd_not_ingested", "company": company, "role_title": role_title}

    if not isinstance(gaps, list) or not gaps:
        return {"ok": False, "error": "invalid_gaps", "detail": "gaps must be a non-empty list"}

    valid_categories = {"missing", "understated", "mismatch"}
    seen_ids = set()
    for i, gap in enumerate(gaps):
        if not isinstance(gap, dict):
            return {"ok": False, "error": "invalid_gap_entry", "index": i}
        gap_id = gap.get("gap_id")
        category = gap.get("category")
        if not gap_id or gap_id in seen_ids:
            return {"ok": False, "error": "invalid_gap_entry", "index": i, "field": "gap_id"}
        seen_ids.add(gap_id)
        if category not in valid_categories:
            return {"ok": False, "error": "invalid_gap_entry", "index": i, "field": "category", "value": category}
        for field in ("jd_criterion", "affected_cv_section", "recommendation"):
            if not gap.get(field):
                return {"ok": False, "error": "invalid_gap_entry", "index": i, "field": field}
        excerpt = gap.get("current_text_excerpt")
        if category == "missing" and excerpt is not None:
            return {
                "ok": False, "error": "invalid_gap_entry", "index": i,
                "field": "current_text_excerpt", "detail": "must be null when category is 'missing'",
            }
        if category in ("understated", "mismatch") and not excerpt:
            return {
                "ok": False, "error": "invalid_gap_entry", "index": i,
                "field": "current_text_excerpt",
                "detail": "required (verbatim CV text) when category is 'understated' or 'mismatch'",
            }

    company_dir = _resolve_company_folder(company, role_title, tracker)
    company_dir.mkdir(parents=True, exist_ok=True)

    lines = [f"# GAP ANALYSIS: {company} — {role_title}", "", f"> Generated: {datetime.now().strftime('%Y-%m-%d')}", ""]
    for gap in gaps:
        lines.append(f"## [{gap['category']}] {gap['jd_criterion']}")
        lines.append(f"- Affected section: {gap['affected_cv_section']}")
        if gap.get("current_text_excerpt"):
            lines.append(f"- Current text: \"{gap['current_text_excerpt']}\"")
        lines.append(f"- Recommendation: {gap['recommendation']}")
        lines.append("")

    output_path = company_dir / "gap_analysis.md"
    output_path.write_text("\n".join(lines), encoding="utf-8")

    # Record output in tracker (tracker already loaded above)
    _record_output(tracker, company, role_title, "gap_analysis", {
        "path": str(output_path),
        "saved_at": _utc_now(),
    })
    _save_tracker(tracker)

    return {
        "ok": True,
        "company": company,
        "role_title": role_title,
        "path": str(output_path),
        "gaps_saved": len(gaps),
    }


@mcp.tool()
def generate_learning_program(company: str, role_title: str) -> dict:
    """Prepare context for generating a skill-gap Learning_Program.

    Uses the missing_skills list from the most recent Match_Score. Returns an
    error if no Match_Score exists yet, or a no-gaps response (no file
    written) if missing_skills is empty.

    Args:
        company: Target employer name.
        role_title: Role title identifying the tracked application.
    """
    tracker = _load_tracker()
    app = _find_application(tracker, company, role_title)
    if app is None:
        return {"ok": False, "error": "jd_not_ingested", "company": company, "role_title": role_title}

    match_score = app.get("match_score")
    if match_score is None:
        return {"ok": False, "error": "no_match_score", "company": company, "role_title": role_title}

    missing_skills = match_score.get("missing_skills", [])
    if not missing_skills:
        return {
            "ok": True,
            "no_gaps": True,
            "company": company,
            "role_title": role_title,
            "message": "No skill gaps were identified; no learning_program.md was created.",
        }

    jd_fields = _load_jd_fields_for_app(app) or {}
    required = {s.lower() for s in (jd_fields.get("required_skills") or [])}
    preferred = {s.lower() for s in (jd_fields.get("preferred_skills") or [])}

    skill_priorities = {}
    for skill in missing_skills:
        low = skill.lower()
        if low in required:
            priority = "high"
        elif low in preferred:
            priority = "medium"
        else:
            priority = "low"
        skill_priorities[skill] = {"priority": priority, "completion_days": LEARNING_PRIORITY_DAYS[priority]}

    return {
        "ok": True,
        "company": company,
        "role_title": role_title,
        "missing_skills": missing_skills,
        "skill_priorities": skill_priorities,
        "instructions": (
            "For each skill in missing_skills, produce a learning-plan entry with: the skill "
            "name, at least one recommended learning resource (course title, certification "
            "name, or documentation URL), an estimated time commitment between 1 and 200 "
            "hours inclusive, the priority level given in skill_priorities for that skill, "
            "and the matching completion_days from skill_priorities. Include at least one "
            "free or publicly available resource per skill in addition to any paid "
            "certification. Where a skill corresponds to a vendor-specific platform (e.g. "
            "Snowflake, Databricks, Salesforce), include the vendor's official certification "
            "pathway as the primary resource and still include at least one free resource as "
            "a secondary entry."
        ),
        "output_schema": {
            "skill": "string",
            "resources": "list of strings (course/cert/doc names or URLs); at least one free",
            "hours": "int 1-200",
            "priority": "high | medium | low",
            "completion_days": "30 | 60 | 90 — must match priority",
        },
        "output_path": str(_resolve_company_folder(company, role_title, tracker) / "learning_program.md"),
        "call_next": "save_learning_program(company, role_title, program)",
    }


@mcp.tool()
def save_learning_program(company: str, role_title: str, program: list) -> dict:
    """Validate and save a Learning_Program to {Company_Folder}/learning_program.md.

    Args:
        company: Target employer name.
        role_title: Role title identifying the tracked application.
        program: List of learning-plan entries (skill, resources, hours,
            priority, completion_days).
    """
    tracker = _load_tracker()
    app = _find_application(tracker, company, role_title)
    if app is None:
        return {"ok": False, "error": "jd_not_ingested", "company": company, "role_title": role_title}

    if not isinstance(program, list) or not program:
        return {"ok": False, "error": "invalid_program", "detail": "program must be a non-empty list"}

    for i, entry in enumerate(program):
        if not isinstance(entry, dict):
            return {"ok": False, "error": "invalid_program_entry", "index": i}
        if not entry.get("skill"):
            return {"ok": False, "error": "invalid_program_entry", "index": i, "field": "skill"}
        resources = entry.get("resources")
        if not isinstance(resources, list) or not resources:
            return {"ok": False, "error": "invalid_program_entry", "index": i, "field": "resources"}
        hours = entry.get("hours")
        if not isinstance(hours, (int, float)) or isinstance(hours, bool) or not (1 <= hours <= 200):
            return {"ok": False, "error": "invalid_program_entry", "index": i, "field": "hours", "value": hours}
        priority = entry.get("priority")
        if priority not in LEARNING_PRIORITY_DAYS:
            return {"ok": False, "error": "invalid_program_entry", "index": i, "field": "priority", "value": priority}
        expected_days = LEARNING_PRIORITY_DAYS[priority]
        if entry.get("completion_days") != expected_days:
            return {
                "ok": False, "error": "invalid_program_entry", "index": i,
                "field": "completion_days", "value": entry.get("completion_days"), "expected": expected_days,
            }

    company_dir = _resolve_company_folder(company, role_title, tracker)
    company_dir.mkdir(parents=True, exist_ok=True)

    priority_order = {"high": 0, "medium": 1, "low": 2}
    sorted_program = sorted(program, key=lambda e: priority_order[e["priority"]])

    lines = [f"# LEARNING PROGRAM: {company} — {role_title}", "", f"> Generated: {datetime.now().strftime('%Y-%m-%d')}", ""]
    for entry in sorted_program:
        lines.append(f"## {entry['skill']} ({entry['priority']} priority, {entry['completion_days']}-day target)")
        lines.append(f"- Estimated time: {entry['hours']} hours")
        lines.append("- Resources:")
        for resource in entry["resources"]:
            lines.append(f"  - {resource}")
        lines.append("")

    output_path = company_dir / "learning_program.md"
    output_path.write_text("\n".join(lines), encoding="utf-8")

    # Record output in tracker (tracker already loaded above)
    _record_output(tracker, company, role_title, "learning_program", {
        "path": str(output_path),
        "saved_at": _utc_now(),
    })
    _save_tracker(tracker)

    return {
        "ok": True,
        "company": company,
        "role_title": role_title,
        "path": str(output_path),
        "entries_saved": len(sorted_program),
    }


@mcp.tool()
def company_research(company: str, focus: str = "general") -> dict:
    """Get a structured research template for a target company, or view existing research.

    Returns a research template with sections covering company background, financials,
    strategy, management, competitors, and interview insights. Use the deep-research
    skill to fill in each section, then save the results to research.md.

    If research.md already exists, returns its content for reference.

    Args:
        company: Target employer name (e.g., "Gartner").
        focus: Area to emphasize in research (e.g., "AI strategy", "public sector",
               "competitors"). Default: "general".
    """
    company_dir = ARTEFACTS_DIR / company
    if not company_dir.exists():
        return {"error": f"Company folder '{company}' not found. Call create_application first."}

    research_file = company_dir / "research.md"
    existing_research = _read_file(research_file)

    template = RESEARCH_TEMPLATE.format(
        company=company,
        date=datetime.now().strftime("%Y-%m-%d"),
        focus=focus,
    )

    return {
        "company": company,
        "research_file": str(research_file),
        "existing_research": existing_research is not None,
        "template": template,
        "next_steps": [
            "1. Use the deep-research skill to research the company",
            "2. Fill in each section of the template",
            "3. Use search_contacts (ai-assistant) to find your contacts at the company",
            "4. Use ask_question (gebiz-awards) for procurement data",
            "5. Save completed research to research.md",
        ] if not existing_research else [
            "Research already exists. Review and update as needed.",
        ],
    }


@mcp.tool()
def save_research(company: str, content: str, focus: str = "general") -> dict:
    """Save company research content to research.md.

    Write the completed research content (typically from the deep-research skill)
    to the company's research.md file. Adds a header with date and focus area.

    Args:
        company: Target employer name (e.g., "Gartner").
        content: The research content in Markdown format.
        focus: Research focus area. Default: "general".
    """
    company_dir = ARTEFACTS_DIR / company
    if not company_dir.exists():
        return {"error": f"Company folder '{company}' not found. Call create_application first."}

    research_file = company_dir / "research.md"
    header = f"# {company} — Company Research\n\n> Generated: {datetime.now().strftime('%Y-%m-%d')} | Focus: {focus}\n\n"
    research_file.write_text(header + content, encoding="utf-8")

    # Record output in tracker
    tracker = _load_tracker()
    _record_output(tracker, company, None, "research", {
        "path": str(research_file),
        "saved_at": _utc_now(),
    })
    _save_tracker(tracker)

    return {
        "company": company,
        "path": str(research_file),
        "content_length": len(content),
        "saved": True,
    }


@mcp.tool()
def map_territory(company: str, accounts: list[str]) -> dict:
    """Get a territory mapping template for specific accounts at a target company.

    Returns a template for mapping your contacts to the target company's key accounts.
    For each account, use the ai-assistant search_contacts tool and sgdi_query tool
    to find and enrich contacts, then save results with save_territory_map.

    Args:
        company: Target employer name (e.g., "Gartner"). Used for the file path.
        accounts: Specific account names to search for contacts
                  (e.g., ["MTI", "GovTech", "MAS"]). These are the organizations
                  where you want to find your contacts.
    """
    company_dir = ARTEFACTS_DIR / company
    if not company_dir.exists():
        return {"error": f"Company folder '{company}' not found. Call create_application first."}

    territory_file = company_dir / "territory_map.md"
    existing = _read_file(territory_file)

    # Build per-account instructions
    account_instructions = []
    for account in accounts:
        account_instructions.append({
            "account": account,
            "mcp_calls": [
                f'ai-assistant search_contacts(company="{account}")',
                f'ai-assistant get_contacts_by_org(organization="{account}")',
                f'sgdi sgdi_query(organisation="{account.lower()}")',
                f'gebiz-awards ask_question(question="What contracts does {account} have with {company}?")',
            ],
        })

    return {
        "company": company,
        "territory_file": str(territory_file),
        "existing_territory_map": existing is not None,
        "accounts": accounts,
        "account_instructions": account_instructions,
        "template": TERRITORY_TEMPLATE.format(
            company=company,
            date=datetime.now().strftime("%Y-%m-%d"),
            account=", ".join(accounts),
        ),
        "next_steps": [
            f"1. For each account ({', '.join(accounts)}), call the MCP tools listed above",
            "2. Gather contact data: name, role, email, phone, organization",
            "3. Note warm paths and account intelligence",
            "4. Call save_territory_map with the formatted content",
        ],
    }


@mcp.tool()
def save_territory_map(company: str, content: str) -> dict:
    """Save territory and contact mapping to territory_map.md.

    Formats and saves the contact data gathered from ai-assistant,
    contacts, sgdi, and gebiz-awards into a structured territory map.

    Args:
        company: Target employer name (e.g., "Gartner").
        content: The territory map content in Markdown format.
    """
    company_dir = ARTEFACTS_DIR / company
    if not company_dir.exists():
        return {"error": f"Company folder '{company}' not found. Call create_application first."}

    territory_file = company_dir / "territory_map.md"
    header = f"# {company} — Territory & Contact Map\n\n> Generated: {datetime.now().strftime('%Y-%m-%d')}\n\n"
    territory_file.write_text(header + content, encoding="utf-8")

    # Record output in tracker
    tracker = _load_tracker()
    _record_output(tracker, company, None, "territory_map", {
        "path": str(territory_file),
        "saved_at": _utc_now(),
    })
    _save_tracker(tracker)

    return {
        "company": company,
        "path": str(territory_file),
        "content_length": len(content),
        "saved": True,
    }


@mcp.tool()
def generate_cover_letter(company: str, tone: str = "storyteller") -> dict:
    """Prepare context for generating a tailored cover letter.

    Reads the JD, research, and CV files from the company folder and returns
    them organized for cover letter generation. The actual content is generated
    by Claude Code using this context.

    Args:
        company: Target employer name (e.g., "Gartner").
        tone: Writing tone: "bold", "conservative", or "storyteller" (default).
              Storyteller uses the user's established "Ground Truth" framing.
    """
    if tone not in COVER_LETTER_TONES:
        return {
            "error": "invalid_tone",
            "message": f"'{tone}' is not a valid tone.",
            "valid_tones": sorted(COVER_LETTER_TONES),
        }

    company_dir = ARTEFACTS_DIR / company
    if not company_dir.exists():
        return {"error": f"Company folder '{company}' not found. Call create_application first."}

    # Read source files
    jd_content = _read_file(company_dir / "JD.md")
    research_content = _read_file(company_dir / "research.md")
    cv_file = _find_cv(company_dir)

    cv_content = None
    if cv_file:
        if cv_file.suffix.lower() == ".pdf":
            cv_content = _extract_pdf_text(cv_file)
        else:
            cv_content = _read_file(cv_file)

    # Find existing cover letter
    existing_letter = None
    for f in company_dir.iterdir():
        if f.is_file() and ("cover" in f.name.lower() or "letter" in f.name.lower()):
            existing_letter = _read_file(f)
            break

    tone_description = COVER_LETTER_TONES[tone]

    missing = []
    if not jd_content:
        missing.append("JD.md — call create_application first")
    if not research_content:
        missing.append("research.md — call company_research first")
    if not cv_content:
        missing.append("CV file — add a CV to the company folder")

    result = {
        "company": company,
        "tone": tone,
        "tone_description": tone_description,
        "source_files": {
            "jd": {"available": jd_content is not None, "length": len(jd_content) if jd_content else 0},
            "research": {"available": research_content is not None, "length": len(research_content) if research_content else 0},
            "cv": {"available": cv_content is not None, "length": len(cv_content) if cv_content else 0},
        },
        "jd_content": jd_content,
        "research_content": research_content,
        "cv_content": cv_content,
        "existing_cover_letter": existing_letter,
        "missing_sources": missing,
        "output_path": str(company_dir / "Cover_Letter.md"),
        "instructions": (
            "Generate a cover letter using the provided context. "
            f"Use a {tone} tone. "
            "Structure it as: an opening paragraph addressing the hiring team, a body of no more than "
            "three paragraphs covering differentiators that map directly to JD requirements or "
            "responsibilities, and a closing paragraph with an explicit request for a meeting or "
            "interview. Embed at least one specific, quantified achievement from the CV/profile "
            "content; if none exists, use the most senior role title and company name as a fallback "
            "rather than fabricating a number. "
            "Write the cover letter in Markdown format and save to the output_path via save_cover_letter."
        ),
    }
    if not research_content:
        result["warning"] = (
            "No company research found (research.md missing) — generating cover letter "
            "using JD and CV context only."
        )
    return result


@mcp.tool()
def save_cover_letter(company: str, content: str, tone: str = "storyteller") -> dict:
    """Save a cover letter to the company folder.

    Requirement 7.5: if a Cover_Letter.md already exists, it is renamed to
    Cover_Letter_v{N}.md (N = existing versioned backups + 1) before the new
    one is written — no backup is made the first time a letter is saved.

    Args:
        company: Target employer name.
        content: The cover letter content in Markdown format.
        tone: The tone used (for the file header).
    """
    company_dir = ARTEFACTS_DIR / company
    if not company_dir.exists():
        return {"error": f"Company folder '{company}' not found."}

    output_path = company_dir / "Cover_Letter.md"
    backed_up_to = None
    if output_path.exists():
        existing_versions = list(company_dir.glob("Cover_Letter_v*.md"))
        backup_path = company_dir / f"Cover_Letter_v{len(existing_versions) + 1}.md"
        output_path.rename(backup_path)
        backed_up_to = str(backup_path)

    header = f"# COVER LETTER: {company.upper()}\n\n> Tone: {tone} | Generated: {datetime.now().strftime('%Y-%m-%d')}\n\n"
    output_path.write_text(header + content, encoding="utf-8")

    # Record output in tracker
    version = len(list(company_dir.glob("Cover_Letter_v*.md"))) + 1  # current version (backup just moved, so +1)
    tracker = _load_tracker()
    _record_output(tracker, company, None, "cover_letter", {
        "path": str(output_path),
        "saved_at": _utc_now(),
        "version": version,
    })
    _save_tracker(tracker)

    return {
        "company": company,
        "path": str(output_path),
        "backed_up_previous_to": backed_up_to,
        "content_length": len(content),
        "saved": True,
    }


@mcp.tool()
def generate_pitch(company: str, format: str = "narrative") -> dict:
    """Prepare context for generating an interview pitch and questions.

    Reads research and territory map files and returns them organized for pitch
    generation. Also checks for MEDDPICC gaps via ai-crm if an opportunity exists.

    Args:
        company: Target employer name (e.g., "Gartner").
        format: Output format: "narrative", "bullet_points", or "star_stories" (default: "narrative").
    """
    company_dir = ARTEFACTS_DIR / company
    if not company_dir.exists():
        return {"error": f"Company folder '{company}' not found. Call create_application first."}

    # Read source files
    jd_content = _read_file(company_dir / "JD.md")
    research_content = _read_file(company_dir / "research.md")
    territory_content = _read_file(company_dir / "territory_map.md")
    cv_file = _find_cv(company_dir)
    cv_content = None
    if cv_file:
        cv_content = _extract_pdf_text(cv_file) if cv_file.suffix.lower() == ".pdf" else _read_file(cv_file)

    existing_pitch = _read_file(company_dir / "pitch.md")

    format_description = PITCH_FORMATS.get(format, PITCH_FORMATS["narrative"])

    missing = []
    if not research_content:
        missing.append("research.md — call company_research first")
    if not territory_content:
        missing.append("territory_map.md — call map_territory first")

    return {
        "company": company,
        "format": format,
        "format_description": format_description,
        "source_files": {
            "jd": {"available": jd_content is not None, "length": len(jd_content) if jd_content else 0},
            "research": {"available": research_content is not None, "length": len(research_content) if research_content else 0},
            "territory_map": {"available": territory_content is not None, "length": len(territory_content) if territory_content else 0},
            "cv": {"available": cv_content is not None, "length": len(cv_content) if cv_content else 0},
        },
        "jd_content": jd_content,
        "research_content": research_content,
        "territory_map_content": territory_content,
        "cv_content": cv_content,
        "existing_pitch": existing_pitch,
        "missing_sources": missing,
        "output_path": str(company_dir / "pitch.md"),
        "instructions": (
            f"Generate an interview pitch using a {format} format. "
            "Include: 1) Your key messages, 2) Questions to ask the interviewers, "
            "3) How to address potential concerns, 4) Stories/evidence from your experience. "
            "After generating, also call ai-crm get_meddpicc_gaps to identify information gaps "
            "for this opportunity. Save the pitch to the output_path."
        ),
    }


@mcp.tool()
def save_pitch(company: str, content: str, format: str = "narrative") -> dict:
    """Save an interview pitch to the company folder.

    Args:
        company: Target employer name.
        content: The pitch content in Markdown format.
        format: The format used (for the file header).
    """
    company_dir = ARTEFACTS_DIR / company
    if not company_dir.exists():
        return {"error": f"Company folder '{company}' not found."}

    pitch_file = company_dir / "pitch.md"
    header = f"# {company} — Interview Pitch\n\n> Format: {format} | Generated: {datetime.now().strftime('%Y-%m-%d')}\n\n"
    pitch_file.write_text(header + content, encoding="utf-8")

    # Record output in tracker
    tracker = _load_tracker()
    _record_output(tracker, company, None, "pitch", {
        "path": str(pitch_file),
        "saved_at": _utc_now(),
    })
    _save_tracker(tracker)

    return {
        "company": company,
        "path": str(pitch_file),
        "content_length": len(content),
        "saved": True,
    }


@mcp.tool()
def tailor_cv(company: str) -> dict:
    """Prepare context for tailoring a CV to a specific job description.

    Reads the JD, research, and base CV files and returns them organized for
    CV tailoring. The actual tailoring is done by Claude Code using this context.

    Args:
        company: Target employer name (e.g., "Gartner").
    """
    company_dir = ARTEFACTS_DIR / company
    if not company_dir.exists():
        return {"error": f"Company folder '{company}' not found. Call create_application first."}

    # Read source files
    jd_content = _read_file(company_dir / "JD.md")
    research_content = _read_file(company_dir / "research.md")
    cv_file = _find_cv(company_dir)
    cv_content = None
    if cv_file:
        cv_content = _extract_pdf_text(cv_file) if cv_file.suffix.lower() == ".pdf" else _read_file(cv_file)

    existing_tailored = _read_file(company_dir / "CV_tailored.md")

    # Requirement 6.1/6.3: pull top strengths from the most recent Match_Score
    # and the role title, if this application is tracked.
    tracker = _load_tracker()
    app = _find_application(tracker, company, None)
    match_score_strengths = []
    role_title = None
    if app:
        role_title = app.get("role_title")
        match_score = app.get("match_score")
        if match_score:
            match_score_strengths = match_score.get("strengths", [])[:3]

    gap_analysis_content = _read_file(company_dir / "gap_analysis.md")

    missing = []
    if not jd_content:
        missing.append("JD.md — call create_application first")
    if not cv_content:
        missing.append("CV file — add a CV (PDF or Markdown) to the company folder")

    instructions = (
        "Generate a tailored CV that highlights experience relevant to the job description. "
        "Keep the same structure as the original CV but: "
        "1) Reorder bullet points to lead with JD-relevant achievements, "
        "2) Emphasize skills and experiences that match key JD requirements, "
        "3) Prepend a tailored professional summary of no more than 120 words naming the target "
        f"company{f', the role \"{role_title}\"' if role_title else ''}, and the top strengths in "
        "match_score_strengths below. "
        "Do NOT fabricate experience — only reorganize and emphasize what's already there. "
        "Do NOT alter any text segment containing a number with a currency symbol or percent sign, "
        "or the words ARR, quota, deal, or target — copy those segments verbatim from the Base_CV. "
        "Save the result to the output_path via save_tailored_cv, passing a diff_summary list of "
        "{section, change_type, description} entries (change_type one of: reorder, condense, add, "
        "remove, replace) describing every change made relative to the Base_CV."
    )
    if gap_analysis_content:
        instructions += " A gap_analysis_content is available below — incorporate its accepted edit recommendations."
    else:
        instructions += " No gap_analysis.md exists for this application — proceed with the Base_CV only, and note its absence."

    return {
        "company": company,
        "source_files": {
            "jd": {"available": jd_content is not None, "length": len(jd_content) if jd_content else 0},
            "research": {"available": research_content is not None, "length": len(research_content) if research_content else 0},
            "cv": {
                "available": cv_content is not None,
                "length": len(cv_content) if cv_content else 0,
                "source_file": str(cv_file) if cv_file else None,
            },
            "gap_analysis": {"available": gap_analysis_content is not None},
        },
        "jd_content": jd_content,
        "research_content": research_content,
        "cv_content": cv_content,
        "existing_tailored_cv": existing_tailored,
        "match_score_strengths": match_score_strengths,
        "gap_analysis_content": gap_analysis_content,
        "gap_analysis_available": gap_analysis_content is not None,
        "missing_sources": missing,
        "output_path": str(company_dir / "CV_tailored.md"),
        "instructions": instructions,
    }


@mcp.tool()
def save_tailored_cv(company: str, content: str, diff_summary: list | None = None) -> dict:
    """Save a tailored CV to the company folder.

    Args:
        company: Target employer name.
        content: The tailored CV content in Markdown format.
        diff_summary: Optional list of {section, change_type, description} entries
            describing changes relative to the Base_CV. Auto-generated via a
            structural diff if omitted.
    """
    company_dir = ARTEFACTS_DIR / company
    if not company_dir.exists():
        return {"error": f"Company folder '{company}' not found."}

    base_cv_content = _resolve_base_cv_content(company_dir)
    if base_cv_content:
        altered = [line for line in _protected_lines(base_cv_content) if line not in content]
        if altered:
            return {
                "error": "fabrication_detected",
                "message": "One or more protected figures from the Base_CV were altered or removed.",
                "altered_segments": altered,
            }

    if diff_summary is not None:
        for i, entry in enumerate(diff_summary):
            if not isinstance(entry, dict) or not entry.get("section") or not entry.get("description"):
                return {"error": "invalid_diff_entry", "index": i, "field": "section/description"}
            if entry.get("change_type") not in VALID_DIFF_CHANGE_TYPES:
                return {
                    "error": "invalid_diff_entry",
                    "index": i,
                    "field": "change_type",
                    "valid_change_types": sorted(VALID_DIFF_CHANGE_TYPES),
                }
    elif base_cv_content:
        diff_summary = _generate_diff_summary(base_cv_content, content)
    else:
        diff_summary = []

    cv_file = company_dir / "CV_tailored.md"
    header = f"# LEE GEE SIN — Tailored CV for {company}\n\n> Generated: {datetime.now().strftime('%Y-%m-%d')}\n\n"
    cv_file.write_text(header + content, encoding="utf-8")

    diff_path = company_dir / "cv_diff_summary.md"
    diff_lines = [f"[{e['section']}] {e['change_type']}: {e['description']}" for e in diff_summary]
    diff_body = "\n".join(diff_lines) if diff_lines else "(no changes recorded)"
    diff_path.write_text(f"# CV Diff Summary: {company}\n\n{diff_body}\n", encoding="utf-8")

    # Record output in tracker
    tracker = _load_tracker()
    _record_output(tracker, company, None, "tailored_cv", {
        "path": str(cv_file),
        "saved_at": _utc_now(),
    })
    _save_tracker(tracker)

    # Create CV record via CVVersioningService
    cv_service = _get_or_create_cv_service()
    try:
        cv_record = cv_service.create_draft_record(
            application_id=company,
            content=content,
            evidence_used=[]
        )
        cv_id = cv_record.cv_id
        version = cv_record.version
        status = cv_record.status.value
    except Exception as e:
        cv_id = None
        version = "unknown"
        status = "error"

    return {
        "company": company,
        "path": str(cv_file),
        "diff_summary_path": str(diff_path),
        "diff_entries": len(diff_summary),
        "content_length": len(content),
        "saved": True,
        "cv_id": cv_id,
        "version": version,
        "status": status,
    }


@mcp.tool()
def save_interview_notes(
    company: str,
    content: str,
    section: str | None = None,
    role_title: str | None = None,
) -> dict:
    """Save or append interview notes for a company's application.

    Creates interview_notes.md if it doesn't exist, or appends under a
    timestamped heading if it does. Optionally includes a section title
    in the heading.

    Args:
        company: Target employer name.
        content: The interview notes content in Markdown format.
        section: Optional section title (e.g., "Recruiter call", "Round 2 feedback").
        role_title: Role title to disambiguate multi-role companies.
    """
    tracker = _load_tracker()
    company_dir = _resolve_company_folder(company, role_title, tracker)
    company_dir.mkdir(parents=True, exist_ok=True)

    notes_path = company_dir / "interview_notes.md"
    timestamp = _utc_now()
    section_title = f" — {section}" if section else ""

    if notes_path.exists():
        # Append under a new timestamped heading
        existing = notes_path.read_text(encoding="utf-8")
        new_entry = f"\n\n---\n\n## {timestamp}{section_title}\n\n{content}\n"
        notes_path.write_text(existing + new_entry, encoding="utf-8")
        appended = True
    else:
        # Create new file with header
        header = f"# {company} — Interview Notes\n\n> Created: {datetime.now().strftime('%Y-%m-%d')}\n\n"
        heading = f"## {timestamp}{section_title}\n\n"
        notes_path.write_text(header + heading + content + "\n", encoding="utf-8")
        appended = False

    # Record output in tracker
    _record_output(tracker, company, role_title, "interview_notes", {
        "path": str(notes_path),
        "saved_at": timestamp,
        "section": section,
        "appended": appended,
    })
    _save_tracker(tracker)

    return {
        "ok": True,
        "company": company,
        "role_title": role_title,
        "path": str(notes_path),
        "appended": appended,
        "content_length": len(content),
    }


@mcp.tool()
def mark_submitted(
    company: str,
    role_title: str | None = None,
    document_types: list[str] | None = None,
) -> dict:
    """Mark CV and/or cover letter as submitted by copying to a submitted/ subfolder.

    Copies the current tailored CV and/or cover letter into a 'submitted' subfolder
    inside the company (or role) directory, creating a snapshot of the exact versions
    that were sent to the employer. Records the submission in the tracker.

    Args:
        company: Target employer name.
        role_title: Role title to disambiguate multi-role companies.
        document_types: Which documents to submit. Default: ["cv", "cover_letter"].
            Valid values: "cv", "cover_letter".
    """
    VALID_DOC_TYPES = {"cv", "cover_letter"}
    if document_types is None:
        document_types = ["cv", "cover_letter"]

    invalid = [dt for dt in document_types if dt not in VALID_DOC_TYPES]
    if invalid:
        return {"ok": False, "error": "invalid_document_type", "invalid": invalid,
                "valid_types": sorted(VALID_DOC_TYPES)}

    tracker = _load_tracker()
    app = _find_application(tracker, company, role_title)
    if app is None:
        # Try to resolve without tracker for legacy folders
        company_dir = ARTEFACTS_DIR / company
        if not company_dir.exists():
            return {"ok": False, "error": "application_not_found",
                    "company": company, "role_title": role_title}
    else:
        company_dir = _resolve_company_folder(company, role_title, tracker)

    submitted_dir = company_dir / "submitted"
    submitted_dir.mkdir(parents=True, exist_ok=True)

    source_map = {
        "cv": company_dir / "CV_tailored.md",
        "cover_letter": company_dir / "Cover_Letter.md",
    }

    files_copied = []
    errors = []

    for doc_type in document_types:
        source = source_map[doc_type]
        if not source.exists():
            errors.append({
                "document_type": doc_type,
                "error": "source_file_missing",
                "path": str(source),
            })
            continue

        if doc_type == "cv":
            dest = submitted_dir / "CV_tailored.md"
        elif doc_type == "cover_letter":
            dest = submitted_dir / "Cover_Letter.md"

        shutil.copy2(source, dest)
        submitted_at = _utc_now()

        if app is not None:
            app.setdefault("submitted", {})[doc_type] = {
                "path": str(dest),
                "submitted_at": submitted_at,
            }

        files_copied.append({
            "document_type": doc_type,
            "source": str(source),
            "dest": str(dest),
            "submitted_at": submitted_at,
        })

    if app is not None:
        _save_tracker(tracker)

    if errors and not files_copied:
        return {"ok": False, "error": "all_sources_missing", "details": errors}

    result = {
        "ok": True,
        "company": company,
        "role_title": role_title,
        "files_copied": files_copied,
        "submitted": app.get("submitted", {}) if app else {},
    }
    if errors:
        result["warnings"] = errors
    return result


@mcp.tool()
def export_document(company: str, document_type: str, format: str, role_title: str | None = None) -> dict:
    """Export a tailored CV or cover letter to PDF or DOCX (Requirement 8).

    Args:
        company: Target employer name.
        document_type: Either "tailored_cv" (CV_tailored.md) or "cover_letter" (Cover_Letter.md).
        format: Either "pdf" or "docx".
        role_title: Optional role title, required to disambiguate companies
            with multiple tracked roles.
    """
    if document_type not in DOCUMENT_TYPE_FILES:
        return {
            "ok": False,
            "error": "invalid_document_type",
            "valid_document_types": sorted(DOCUMENT_TYPE_FILES),
        }
    if format not in EXPORT_FORMATS:
        return {"ok": False, "error": "invalid_format", "valid_formats": sorted(EXPORT_FORMATS)}

    tracker = _load_tracker()
    try:
        company_dir = _resolve_company_folder(company, role_title, tracker)
    except AmbiguousRoleError as e:
        return {"ok": False, "error": "ambiguous_role", "company": e.company, "roles": e.roles}

    source_file = company_dir / DOCUMENT_TYPE_FILES[document_type]
    if not source_file.exists():
        return {
            "ok": False,
            "error": "source_file_missing",
            "message": f"{source_file.name} not found for {company}. Generate it first.",
        }

    missing_deps = _check_export_deps(format)
    if missing_deps:
        return {"ok": False, "error": "missing_export_dep", "missing_dependencies": missing_deps}

    date_str = datetime.now().strftime("%Y-%m-%d")
    output_path = company_dir / f"{document_type}_{company}_{date_str}.{format}"

    try:
        if format == "pdf":
            _export_to_pdf(source_file, output_path)
        else:
            _export_to_docx(source_file, output_path)
    except Exception as e:
        return {"ok": False, "error": "export_failed", "message": str(e)}

    return {
        "ok": True,
        "company": company,
        "role_title": role_title,
        "document_type": document_type,
        "format": format,
        "output_path": str(output_path),
    }


# ---------------------------------------------------------------------------
# Daily Discovery tools (LinkedIn job digest integration)
# ---------------------------------------------------------------------------

@mcp.tool()
def review_daily_discoveries(date: str = "") -> dict:
    """Review daily job discoveries from the LinkedIn email digest.

    Returns the day's curated job cards surfaced by the 2am digest job.
    Use this each morning to see new job opportunities, then use
    score_match on individual jobs, and ingest_from_discovery to add
    them to the tracker.

    Args:
        date: ISO date string (e.g., "2026-08-06"). Defaults to today.
    """
    from datetime import date as _date

    if not date:
        date = _date.today().isoformat()

    digest_dir = ARTEFACTS_DIR / "digests"
    digest_path = digest_dir / f"{date}.md"

    if not digest_path.exists():
        return {"ok": False, "error": "no_digest_found", "date": date,
                "hint": f"No digest file found for {date}. The job digest script may not have run yet."}

    content = digest_path.read_text(encoding="utf-8")

    # Parse the markdown into structured data
    jobs = []
    current_job = None
    in_below_threshold = False

    for line in content.split("\n"):
        # Section boundary: "## Below Threshold" marks the start of the
        # below-threshold section — stop collecting surfaced jobs.
        if line.strip().startswith("## Below Threshold"):
            in_below_threshold = True
            current_job = None
            continue
        # New section header resets below-threshold flag
        if line.strip().startswith("## ") and not line.strip().startswith("## Below"):
            in_below_threshold = False
            current_job = None

        if line.startswith("### "):
            header = line.lstrip("# ").strip()
            if " — " in header:
                company, title = header.split(" — ", 1)
            else:
                company, title = "", header
            category = "below_threshold" if in_below_threshold else "surfaced"
            current_job = {"company": company.strip(), "title": title.strip(), "category": category}
            jobs.append(current_job)
        elif current_job is not None:
            line = line.strip()
            if line.startswith("- **Location:**"):
                current_job["location"] = line.replace("- **Location:**", "").strip()
            elif line.startswith("- **URL:**"):
                current_job["url"] = line.replace("- **URL:**", "").strip()
            elif line.startswith("- **Snippet:**"):
                current_job["snippet"] = line.replace("- **Snippet:**", "").strip()
            elif line.startswith("- **Category:**"):
                current_job["category"] = line.replace("- **Category:**", "").strip()
            elif line.startswith("- **Reason:**"):
                current_job["reason"] = line.replace("- **Reason:**", "").strip()

    surfaced = [j for j in jobs if j.get("category") != "below_threshold"]
    below_threshold = [j for j in jobs if j.get("category") == "below_threshold"]

    return {
        "ok": True,
        "date": date,
        "surfaced": surfaced,
        "below_threshold": below_threshold,
        "total_surfaced": len(surfaced),
        "total_below_threshold": len(below_threshold),
        "hint": "Use score_match on individual jobs for full LLM-based matching, then ingest_from_discovery to add to tracker.",
    }


@mcp.tool()
def ingest_from_discovery(company: str, date: str) -> dict:
    """Ingest a discovered job from the daily digest into the tracker.

    Finds the job in the digest file for the given date, creates a tracker
    entry (stage=new), and writes JD.md to the company folder.

    Args:
        company: Company name (must match a job in the digest for that date).
        date: ISO date string matching the digest file (e.g., "2026-08-06").
    """
    digest_dir = ARTEFACTS_DIR / "digests"
    digest_path = digest_dir / f"{date}.md"

    if not digest_path.exists():
        return {"ok": False, "error": "no_digest_found", "date": date}

    content = digest_path.read_text(encoding="utf-8")

    # Parse to find the matching job
    current_job = None
    jobs = []
    for line in content.split("\n"):
        if line.startswith("### "):
            header = line.lstrip("# ").strip()
            if " — " in header:
                comp, title = header.split(" — ", 1)
            else:
                comp, title = "", header
            current_job = {"company": comp.strip(), "title": title.strip()}
            jobs.append(current_job)
        elif current_job is not None:
            line = line.strip()
            if line.startswith("- **Location:**"):
                current_job["location"] = line.replace("- **Location:**", "").strip()
            elif line.startswith("- **URL:**"):
                current_job["url"] = line.replace("- **URL:**", "").strip()
            elif line.startswith("- **Snippet:**"):
                current_job["snippet"] = line.replace("- **Snippet:**", "").strip()

    # Find matching job (case-insensitive company match)
    company_lower = company.lower()
    match = None
    for job in jobs:
        if job["company"].lower() == company_lower:
            match = job
            break

    if not match:
        available = [j["company"] for j in jobs]
        return {"ok": False, "error": "job_not_found", "company": company,
                "available_companies": available}

    # Check if already in tracker
    tracker = _load_tracker()
    for app in tracker.get("applications", []):
        if app.get("company", "").lower() == company_lower:
            return {"ok": False, "error": "already_tracked", "company": company,
                    "existing_id": app.get("id"), "existing_stage": app.get("stage")}

    # Create company folder and JD.md
    try:
        company_dir = _resolve_company_folder(match["company"], match["title"], tracker)
    except AmbiguousRoleError as e:
        return {"ok": False, "error": "ambiguous_role", "company": e.company, "roles": e.roles}
    company_dir.mkdir(parents=True, exist_ok=True)

    # Build JD content from digest data
    jd_lines = [
        f"# {match['title']}",
        "",
        f"**Company:** {match['company']}",
        f"**Location:** {match.get('location', 'Not specified')}",
        f"**Source:** LinkedIn job discovery ({date})",
    ]
    if match.get("url"):
        jd_lines.append(f"**URL:** {match['url']}")
    jd_lines.append("")
    jd_lines.append(match.get("snippet", "No description available."))
    jd_lines.append("")
    jd_content = "\n".join(jd_lines)

    jd_path = company_dir / "JD.md"
    jd_path.write_text(jd_content, encoding="utf-8")

    # Create tracker entry using existing patterns
    now = _utc_now()
    new_app = _create_application_record(match["company"], match["title"], str(jd_path))
    if match.get("url"):
        new_app["jd_source_url"] = match["url"]
    tracker["applications"].append(new_app)
    _save_tracker(tracker)

    return {
        "ok": True,
        "application_id": new_app["id"],
        "company": match["company"],
        "role_title": match["title"],
        "stage": "new",
        "folder_path": str(company_dir),
        "jd_path": str(jd_path),
        "hint": "Use score_match for full LLM-based matching, then save_match_score to record results.",
    }


if __name__ == "__main__":
    mcp.run(transport="streamable-http" if MCP_MODE == "http" else "stdio")
