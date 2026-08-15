#!/usr/bin/env python3
"""Generate Word documents for DXC job application."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = "/Users/gslee/Projects/Job-Applications/DXC"

# ============================================================
# 1. CV Document
# ============================================================
def create_cv():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    
    # --- Header ---
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run('LEE GEE SIN')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run('Singapore | 98212429 | geesin.lee@gmail.com | linkedin.com/in/gee-sin-lee-7286b7')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # --- Title ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CLIENT PARTNER | PUBLIC SECTOR | IT SERVICES')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # --- Professional Summary ---
    doc.add_paragraph()
    summary = doc.add_paragraph()
    run = summary.add_run(
        'Enterprise Sales & Client Partner executive with 20+ years of consultative solutioning and high-value '
        'account management in the Singapore Public Sector. Proven track record of owning strategic government accounts, '
        'driving revenue growth through upsell/cross-sell of end-to-end IT services, and delivering complex digital '
        'transformation programmes from pilot to production. Deep relationships across MTI, MAS, GovTech, IMDA, '
        'Synapxe, JTC, and other government agencies. Strong financial acumen in account planning, pipeline management, '
        'forecasting, and margin optimisation.'
    )
    run.font.size = Pt(10)
    
    # --- Core Competencies ---
    doc.add_paragraph()
    heading = doc.add_heading('CORE COMPETENCIES', level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    competencies = [
        ('Public Sector Account Management:', ' MTI, MAS, GovTech, IMDA, Synapxe, JTC, IDA/Smart Nation — long-term relationship building, formalised tendering, and procurement navigation'),
        ('End-to-End IT Services Selling:', ' Application development, managed services, cloud, analytics, cybersecurity (GovTech), and infrastructure (Telcos, network) — from consultative discovery to contract signature'),
        ('Account Farming & Growth:', ' Expanding existing accounts through strategic upsell/cross-sell; 300% pipeline expansion in SG Public Sector'),
        ('Strategic Orchestration:', ' Leading cross-functional squads (Architects, Partners, Bid Management, Delivery) to shape proposals and win business'),
        ('Financial Management:', ' Revenue forecasting, margin tracking, P&L accountability, and account-level profitability'),
        ('AI & Emerging Technology:', ' Agentic AI, Enterprise MCP, iPaaS automation, AI/ML integration — bridging innovation and business outcomes'),
    ]
    
    for bold_text, normal_text in competencies:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bold_text)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(normal_text)
        run.font.size = Pt(10)
    
    # --- Professional Experience ---
    doc.add_paragraph()
    heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Workato
    add_role(doc, 
        'Workato Inc', 'Enterprise Account Executive – SG Public Sector & Malaysia', 'Sep 2024 – Present',
        [
            'Own and manage strategic accounts across Singapore government bodies and Malaysia GLCs, acting as the primary point of contact and trusted advisor for AI-driven automation (iPaaS) adoption.',
            'Achieved 125% of annual quota ($1.2M ARR target), generating $1.5M ARR in revenue — demonstrating consistent over-performance in a territory requiring deep consultative engagement.',
            'Closed the largest deal at $800K ARR, a landmark automation platform engagement with a Singapore government agency, showcasing ability to navigate complex public sector procurement cycles.',
            'Drive upsell and cross-sell opportunities by positioning end-to-end automation solutions — from data integration to Agentic AI orchestration — expanding account footprint within each client.',
            'Architected "automation-first" strategies that serve as the foundation for Agentic AI and Enterprise MCP implementation, aligning technical innovation with policy-led business outcomes.',
            'Partner closely with solutioning and delivery teams to shape proposals, manage pilot-to-production lifecycles, and ensure successful project outcomes for government stakeholders.',
            'Monitor account financials, including revenue tracking, pipeline forecasting, and margin management across the portfolio.',
        ]
    )
    
    # VISEO
    add_role(doc,
        'VISEO SAS', 'Business Development Director – Salesforce & Bespoke Solutions', '2019 – 2025',
        [
            'Achieved 200–400% of annual quota across 3 consecutive years — $6M+ cumulative revenue on combined targets of ~$3.7M (FY2022–FY2024), demonstrating consistent, repeatable over-performance in the Singapore Public Sector.',
            'Closed the largest deal at $5.6M, a landmark bespoke solution engagement with senior government stakeholders at MTI, MAS, GovTech, IMDA, Synapxe, and JTC, alongside a $3.5M expansion deal (FY2023) — showcasing ability to both land new accounts and grow existing ones.',
            'Managed end-to-end sales of SaaS, bespoke software, Salesforce consulting, and managed services for Public Sector and enterprise clients — selling full-lifecycle IT solutions from application development to ongoing managed services.',
            'Orchestrated complex, high-stakes public sector engagements by aligning internal technical squads with senior government stakeholders, delivering bespoke solutions valued at over $15M.',
            'Developed and executed account plans focused on revenue growth and profitability, utilising CRM-driven insights and consultative selling methodologies (MEDDPICC).',
            'Led multiple formalised tendering processes and pilot-to-production cycles, transforming "off-the-shelf" limitations into tailored, high-impact organisational outcomes.',
            'Achieved 300% expansion of strategic client pipelines within the Singapore Public Sector and Malaysia GLCs through data-led market development and relationship cultivation.',
            'Specialised in high-touch, consultative sales where off-the-shelf solutions were insufficient — demonstrating deep understanding of client-specific challenges and co-creating bespoke solutions.',
        ]
    )
    
    # NCS
    add_role(doc,
        'NCS Group', 'Client Partner / Enterprise Sales Lead, Global Business', '2016 – 2018',
        [
            'Owned and grew strategic public sector accounts, acting as the primary relationship manager for Singapore government agencies driving digital transformation (specific agencies under NDA).',
            'Led digital transformation bids for Public Sector agencies, integrating AI/ML into citizen-facing platforms and bridging legacy infrastructure with Smart City capabilities to drive policy-led outcomes.',
            'Drove upsell and cross-sell of end-to-end IT services — application development, managed services, cloud, analytics, cybersecurity, and infrastructure — across assigned accounts.',
            'Partnered with solutioning and delivery teams to shape proposals, win competitive tenders, and oversee end-to-end delivery ensuring projects met contractual commitments and client KPIs.',
            'Maintained high levels of client satisfaction through consistent engagement and successful delivery outcomes across multiple concurrent programmes.',
            'Managed complex, multi-stakeholder engagements across government agencies, navigating formalised public sector procurement and tendering processes (specific agencies under NDA).',
        ]
    )
    
    # ARM UK
    add_role(doc,
        'ARM UK', 'Business Development / Staff Engineer – IoT Platform & Security Solutions', '2018 – 2019',
        [
            'Bridged IT & OT for Telcos and Enterprises, leveraging IoT innovations and cybersecurity solutions to drive business outcomes.',
            'Managed strategic accounts and identified cross-sell opportunities for ARM\'s IoT platform and security product portfolio.',
            'Drove cybersecurity and infrastructure engagements with GovTech, positioning ARM\'s security solutions for government digital infrastructure programmes.',
            'Built deep domain expertise in telecommunications infrastructure and network security, advising Telcos on securing IoT deployments at scale.',
        ]
    )
    
    # Actility
    add_role(doc,
        'Actility SA', 'Business Development Director – Telco Solutions', '2015 – 2016',
        [
            'Acted as Regional Technical Sales Director for APAC, securing major IoT and Smart City projects, including IDA\'s Smart Nation initiative.',
            'Drove account growth through consultative engagement with government and enterprise stakeholders.',
        ]
    )
    
    # DigitalRoute
    add_role(doc,
        'DigitalRoute AB', 'Business Development Director – Telco Solutions', '2011 – 2015',
        [
            'Managed regional technical sales across APAC, opening new markets in Greater China and ANZ.',
            'Won a major project with Jio Reliance, demonstrating ability to penetrate new markets and close large-scale deals.',
            'Sold network infrastructure and data mediation solutions to Tier-1 Telcos across APAC, building deep expertise in telecommunications infrastructure and BSS/OSS platforms.',
        ]
    )
    
    # Comverse
    add_role(doc,
        'Comverse Inc (now Amdocs)', 'Business Development Manager – Telco Solutions', '2008 – 2011',
        [
            'Led a regional sales and presales team in Southeast Asia, achieving the largest global win for a new Mobile Advertising Platform at SingTel.',
            'Delivered network infrastructure and BSS/OSS solutions to major Telcos, building foundational expertise in telecommunications infrastructure sales.',
        ]
    )
    
    # SingTel
    add_role(doc,
        'SingTel & Optus', 'Senior Manager in Group IT – Business Solutions', '2003 – 2005',
        [
            'Managed large-scale transformation projects and led functional requirement development for enterprise IT systems.',
        ]
    )
    
    # Capgemini
    add_role(doc,
        'Capgemini', 'Senior Consultant / Account Manager – Telco & Digital', '2000 – 2003',
        [
            'Led strategic accounts for Tier-1 Telco providers in APAC & Europe, focusing on digital customer experience and BSS/OSS modernisation.',
            'Orchestrated large-scale consulting engagements, moving clients from legacy hardware-centric models to software-defined architectures.',
            'Managed complex delivery ecosystems involving third-party vendors and internal offshore delivery centres.',
            'Delivered network infrastructure and system integration projects for Telco clients, including BSS/OSS transformation and infrastructure modernisation.',
        ]
    )
    
    # Logica
    add_role(doc,
        'Logica (now CGI)', 'Senior Solution Consultant – Telco Solutions', '2005 – 2008',
        [
            'Delivered technical consultancy and business analysis for large-scale enterprise system integrations.',
            'Developed foundational expertise in complex system architecture and project governance.',
        ]
    )
    
    # --- Key Industry Impact ---
    doc.add_paragraph()
    heading = doc.add_heading('KEY INDUSTRY IMPACT', level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    impacts = [
        ('Strategic Stakeholder Management:', ' Orchestrated complex public sector engagements aligning internal technical squads with senior government stakeholders (MTI, MAS, GovTech, IMDA, Synapxe, JTC) to deliver bespoke solutions valued at over $15M.'),
        ('Consultative Tendering & Solutions:', ' Successfully led multiple formalised tendering processes and pilot-to-production cycles, transforming off-the-shelf limitations into tailored, high-impact outcomes.'),
        ('Data-Led Market Expansion:', ' Utilised CRM-driven insights to identify and cultivate new sector relationships, resulting in a 300% expansion of strategic client pipelines within the Singapore Public Sector and Malaysia GLCs.'),
    ]
    
    for bold_text, normal_text in impacts:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bold_text)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(normal_text)
        run.font.size = Pt(10)
    
    # --- Education ---
    doc.add_paragraph()
    heading = doc.add_heading('EDUCATION & CERTIFICATIONS', level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    education = [
        ('FlexiMasters in Artificial Intelligence and AI Ethics', ' — Nanyang Technological University (NTU)\n  Coursework: LLMs, Deep Learning, Responsible AI'),
        ('Master of Technology (Software Engineering)', ' — Institute of Systems Science (ISS), NUS'),
        ('Bachelor of Science (Computer Science)', ' — National University of Singapore (NUS)'),
        ('Specialist Diploma (Business and Big Data Analytics)', ' — Nanyang Polytechnic'),
    ]
    
    for bold_text, normal_text in education:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bold_text)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(normal_text)
        run.font.size = Pt(10)
    
    # --- Technical Skills ---
    doc.add_paragraph()
    heading = doc.add_heading('TECHNICAL SKILLS', level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    skills = [
        ('Cloud/AI:', ' Claude, OpenAI, Google AI Stack, Salesforce, Workato Agentic, MLOps'),
        ('Methodologies:', ' Consultative Selling, MEDDPICC, Agile/Scrum, ITIL, Design Thinking'),
        ('Languages:', ' English (Fluent), Mandarin (Professional Working Proficiency)'),
    ]
    
    for bold_text, normal_text in skills:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bold_text)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(normal_text)
        run.font.size = Pt(10)
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    filepath = os.path.join(OUTPUT_DIR, 'CV LEE Gee Sin 2026 - DXC Client Partner Public Sector.docx')
    doc.save(filepath)
    print(f'CV saved to: {filepath}')


def add_role(doc, company, role, tenure, bullets):
    """Add a role entry to the document."""
    p = doc.add_paragraph()
    run = p.add_run(company)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run = p.add_run(f'  |  {role}  |  {tenure}')
    run.font.size = Pt(10)
    
    for bullet in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        # Bold the first phrase before the first comma or period if it starts with a verb
        run = bp.add_run(bullet)
        run.font.size = Pt(10)


# ============================================================
# 2. DXC Application Form
# ============================================================
def create_form():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    
    # Title
    title = doc.add_heading('Relevant Experience / Skills Summary', level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Sales Deals Details — Give for last 3 companies / 5 to 7 years')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    p = doc.add_paragraph()
    run = p.add_run('Applicant: ')
    run.bold = True
    run = p.add_run('LEE Gee Sin')
    run = p.add_run('\nRole Applied: ')
    run.bold = True
    run = p.add_run('Client Partner – Public Sector (IT Services), Singapore')
    
    # Company entries
    companies = [
        {
            'name': 'Workato Inc',
            'role': 'Enterprise Account Executive – SG Public Sector and Malaysia',
            'tenure': 'Sep 2024 – Present',
            'solutions': 'AI-driven automation (iPaaS), Agentic AI orchestration, Enterprise MCP, automation-first strategies for government bodies; breaking data silos across agencies.',
            'industries': 'Public Sector (Singapore Government / GovTech / GCC), Malaysia GLCs',
            'target': '$1.2M ARR',
            'achievements': '$1.5M ARR (125% quota attainment); largest deal closed at $800K ARR',
            'clients': 'Singapore government agencies (GovTech), Malaysia GLCs',
            'reason': 'Seeking a strategic client partner role with deeper end-to-end IT services ownership and long-term public sector account growth — aligned with DXC\'s mission.',
        },
        {
            'name': 'VISEO SAS',
            'role': 'Business Development Director – Salesforce & Bespoke Solutions',
            'tenure': '2019 – 2025',
            'solutions': 'SaaS, bespoke software development, Salesforce consulting and implementation, managed services; high-touch consultative sales for complex, non-off-the-shelf client requirements.',
            'industries': 'Public Sector (Singapore – MTI, MAS, GovTech, IMDA, Synapxe, JTC), Telecommunications, FSI',
            'target': '$1.5M (FY2024) / $1.2M (FY2023) / $1.0M (FY2022)',
            'achievements': '$6M+ cumulative revenue over 3 years (FY2022–FY2024); 200–400% quota attainment each year; largest deal closed at $5.6M; $3.5M expansion deal (FY2023); total bespoke solutions portfolio valued at over $15M',
            'clients': 'MTI (Ministry of Trade and Industry), MAS (Monetary Authority of Singapore), GovTech, IMDA (Infocomm Media Development Authority), Synapxe, JTC (Jurong Town Corporation), Singapore Public Sector agencies',
            'reason': 'VISEO\'s acquisition/restructuring shifted strategic direction; seeking a role with stronger public sector focus and broader IT services portfolio.',
        },
        {
            'name': 'NCS Group',
            'role': 'Client Partner / Enterprise Sales Lead, Global Business',
            'tenure': '2016 – 2018',
            'solutions': 'Digital transformation solutions for Public Sector agencies; AI/ML integration into citizen-facing platforms; Smart City capabilities; bridging legacy infrastructure to modern architectures.',
            'industries': 'Public Sector (Singapore Government agencies), Smart Nation initiatives',
            'target': 'Not available (role pre-dated current employer\'s quota tracking framework)',
            'achievements': 'Not available (role pre-dated current employer\'s quota tracking framework); key achievements include leading digital transformation bids for multiple Singapore Public Sector agencies, integrating AI/ML into citizen-facing platforms, and bridging legacy infrastructure with Smart City capabilities',
            'clients': 'Singapore Public Sector agencies (specific names under NDA); IDA (Smart Nation initiative)',
            'reason': 'Pursued opportunity to expand into IoT/Smart City regional leadership; role at Actility offered broader APAC scope.',
        },
    ]
    
    for i, company in enumerate(companies, 1):
        doc.add_paragraph()
        heading = doc.add_heading(f'Company {i}', level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        fields = [
            ('Company name:', company['name']),
            ('Role/Tenure:', f"{company['role']} | {company['tenure']}"),
            ('Solutions/Services Sold:', company['solutions']),
            ('Industries:', company['industries']),
            ('Sales Target:', company['target']),
            ('Sales Achievements:', company['achievements']),
            ('Notable clients:', company['clients']),
            ('Reason for leaving:', company['reason']),
        ]
        
        for label, value in fields:
            p = doc.add_paragraph()
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(10)
            run = p.add_run(f' {value}')
            run.font.size = Pt(10)
    
    # Notice Period
    doc.add_paragraph()
    heading = doc.add_heading('Notice Period', level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph()
    run = p.add_run('2 weeks')
    run.font.size = Pt(10)
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    filepath = os.path.join(OUTPUT_DIR, 'DXC - Relevant Experience_Skills Summary - FILLED.docx')
    doc.save(filepath)
    print(f'Form saved to: {filepath}')


if __name__ == '__main__':
    create_cv()
    create_form()
    print('\nDone! Both Word documents created.')